import tkinter as tk
from tkinter import ttk, messagebox, filedialog, simpledialog
import datetime
from datetime import date
import calendar
from decimal import Decimal
import os
import sys
import logging
from app import SalaryApp
from gui_vacation_tab import VacationTab


# Настройка логирования
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('gui.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Проверяем наличие tkcalendar, если его нет - выдаем информативное сообщение
try:
    import tkcalendar
except ImportError:
    messagebox.showerror(
        "Ошибка импорта", 
        "Отсутствует пакет tkcalendar. Установите его командой:\npip install tkcalendar"
    )
    sys.exit(1)


class SalaryCalculatorGUI:
    """Графический интерфейс для системы расчета зарплаты"""
    
    def __init__(self, master):
        """
        Инициализация графического интерфейса
        
        :param master: корневой элемент tkinter
        """
        self.master = master
        self.master.title("Система расчета зарплаты для лицея ПолессГУ")
        self.master.geometry("1200x700")
        self.master.minsize(900, 600)
        
        # Инициализация комбобокса отчетов
        self.report_teacher_combo = None  # Создаем атрибут сразу

        # Настройка стилей
        self.style = ttk.Style()
        self.style.configure("TLabel", font=("Arial", 10))
        self.style.configure("TButton", font=("Arial", 10))
        self.style.configure("TEntry", font=("Arial", 10))
        self.style.configure("Treeview", font=("Arial", 10))
        self.style.configure("Treeview.Heading", font=("Arial", 10, "bold"))
        
        # Конфигурация подключения к базе данных (в реальном приложении должна загружаться из файла)
        self.db_config = {
            'host': 'localhost',
            'database': 'salary_calculator2',
            'user': 'postgres',
            'password': '123321445',
            'port': '5432'
        }
        
        # Текущий выбранный преподаватель
        self.current_teacher_id = None
        
        try:
            # Инициализация приложения
            self.app = SalaryApp(self.db_config)
            
            # Создание основного интерфейса
            self._create_widgets()
            
            # При закрытии окна корректно завершаем работу
            self.master.protocol("WM_DELETE_WINDOW", self._on_close)
            
            # Инициализация начального представления
            self._load_teachers()
            self.update_status("Приложение готово к работе")
        except Exception as e:
            logger.error(f"Ошибка при инициализации приложения: {str(e)}")
            # messagebox.showerror("Ошибка", f"Ошибка при инициализации приложения: {str(e)}")
    
    def _create_widgets(self):
        """Создание элементов интерфейса"""
        # Создание главного меню
        self._create_menu()
        
        # Создание вкладок
        self.notebook = ttk.Notebook(self.master)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Вкладка "Преподаватели"
        self.teachers_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.teachers_frame, text="Преподаватели")
        
        # Вкладка "Расчет зарплаты"
        self.salary_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.salary_frame, text="Расчет зарплаты")
        
        # Вкладка "Отпуска" - временная, будет заменена в _setup_vacation_tab
        self.vacation_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.vacation_frame, text="Отпуска")
        
        # Вкладка "Отчеты"
        self.reports_frame = ttk.Frame(self.notebook)
        self.notebook.add(self.reports_frame, text="Отчеты")
        
        # Настроим каждую вкладку
        self._setup_teachers_tab()
        self._setup_salary_tab()
        self._setup_vacation_tab()
        self._setup_reports_tab()
        
        # Статусная строка
        self.status_var = tk.StringVar()
        self.status_var.set("Готово")
        self.status_bar = ttk.Label(self.master, textvariable=self.status_var, relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.pack(side=tk.BOTTOM, fill=tk.X)
    
    def _create_menu(self):
        """Создание главного меню"""
        menubar = tk.Menu(self.master)
        
        # Меню "Файл"
        file_menu = tk.Menu(menubar, tearoff=0)
        file_menu.add_command(label="Настройки базы данных", command=self._show_db_settings)
        file_menu.add_separator()
        file_menu.add_command(label="Выход", command=self._on_close)
        menubar.add_cascade(label="Файл", menu=file_menu)
        
        # Меню "Преподаватели"
        teachers_menu = tk.Menu(menubar, tearoff=0)
        teachers_menu.add_command(label="Добавить преподавателя", command=self._show_add_teacher_dialog)
        teachers_menu.add_command(label="Импортировать из CSV", command=self._import_teachers_from_csv)
        teachers_menu.add_command(label="Экспортировать в CSV", command=self._export_teachers_to_csv)
        menubar.add_cascade(label="Преподаватели", menu=teachers_menu)
        
        # Меню "Расчеты"
        calc_menu = tk.Menu(menubar, tearoff=0)
        calc_menu.add_command(label="Расчет зарплаты", command=lambda: self.notebook.select(1))
        calc_menu.add_command(label="Расчет отпускных", command=lambda: self.notebook.select(2))
        menubar.add_cascade(label="Расчеты", menu=calc_menu)
        
        # Меню "Отчеты"
        reports_menu = tk.Menu(menubar, tearoff=0)
        reports_menu.add_command(label="Отчет по зарплате", command=lambda: self._generate_salary_report())
        reports_menu.add_command(label="Отчет по отпускам", command=lambda: self.vacation_tab_manager._export_vacation_report() if hasattr(self, 'vacation_tab_manager') else None)
        menubar.add_cascade(label="Отчеты", menu=reports_menu)
        
        # Меню "Справка"
        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(label="Руководство пользователя", command=self._show_user_manual)
        help_menu.add_command(label="О программе", command=self._show_about)
        menubar.add_cascade(label="Справка", menu=help_menu)
        
        self.master.config(menu=menubar)
    
    def _setup_teachers_tab(self):
        """Настройка вкладки 'Преподаватели'"""
        # Фрейм для списка преподавателей
        list_frame = ttk.LabelFrame(self.teachers_frame, text="Список преподавателей")
        list_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Поле поиска
        search_frame = ttk.Frame(list_frame)
        search_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(search_frame, text="Поиск:").pack(side=tk.LEFT, padx=5)
        self.search_var = tk.StringVar()
        search_entry = ttk.Entry(search_frame, textvariable=self.search_var, width=30)
        search_entry.pack(side=tk.LEFT, padx=5)
        search_entry.bind("<KeyRelease>", self._on_search)
        
        # Таблица преподавателей
        columns = ('id', 'name', 'position', 'academic_degree', 'qualification_category', 'hourly_rate', 'experience_years')
        self.teachers_tree = ttk.Treeview(list_frame, columns=columns, show='headings')
        
        # Заголовки столбцов
        self.teachers_tree.heading('id', text='ID')
        self.teachers_tree.heading('name', text='ФИО')
        self.teachers_tree.heading('position', text='Должность')
        self.teachers_tree.heading('academic_degree', text='Уч. степень')
        self.teachers_tree.heading('qualification_category', text='Категория')
        self.teachers_tree.heading('hourly_rate', text='Ставка (руб/ч)')
        self.teachers_tree.heading('experience_years', text='Стаж (лет)')
        
        # Ширина столбцов
        self.teachers_tree.column('id', width=50, anchor=tk.CENTER)
        self.teachers_tree.column('name', width=200)
        self.teachers_tree.column('position', width=150)
        self.teachers_tree.column('academic_degree', width=130)
        self.teachers_tree.column('qualification_category', width=100)
        self.teachers_tree.column('hourly_rate', width=100, anchor=tk.CENTER)
        self.teachers_tree.column('experience_years', width=80, anchor=tk.CENTER)
        
        # Полоса прокрутки
        scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL, command=self.teachers_tree.yview)
        self.teachers_tree.configure(yscroll=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.teachers_tree.pack(fill=tk.BOTH, expand=True)
        
        # Обработчик выбора преподавателя
        self.teachers_tree.bind('<<TreeviewSelect>>', self._on_teacher_select)
        
        # Фрейм для кнопок управления
        btn_frame = ttk.Frame(list_frame)
        btn_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Кнопки
        ttk.Button(btn_frame, text="Добавить", command=self._show_add_teacher_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Редактировать", command=self._show_edit_teacher_dialog).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Удалить", command=self._delete_teacher).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Обновить", command=self._load_teachers).pack(side=tk.RIGHT, padx=5)
        
        # Фрейм для информации о преподавателе
        info_frame = ttk.LabelFrame(self.teachers_frame, text="Информация о преподавателе")
        info_frame = ttk.LabelFrame(self.teachers_frame, text="Информация о преподавателе", width=300)
        info_frame.pack(side=tk.RIGHT, fill=tk.BOTH, padx=5, pady=5, expand=False, ipadx=5, ipady=5)
        
        info_inner_frame = ttk.Frame(info_frame)
        info_inner_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Поля информации
        ttk.Label(info_inner_frame, text="ID:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_id_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_id_var).grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="ФИО:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_name_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_name_var, wraplength=200).grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Должность:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_position_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_position_var).grid(row=2, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Ученая степень:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_degree_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_degree_var).grid(row=3, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Квалификация:").grid(row=4, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_qualification_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_qualification_var).grid(row=4, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Стаж (лет):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_experience_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_experience_var).grid(row=5, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Ставка (руб/час):").grid(row=6, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_rate_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_rate_var).grid(row=6, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Дата приема:").grid(row=7, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_hire_date_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_hire_date_var).grid(row=7, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Дата рождения:").grid(row=8, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_birth_date_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_birth_date_var).grid(row=8, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Молодой спец.:").grid(row=9, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_young_specialist_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_young_specialist_var).grid(row=9, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(info_inner_frame, text="Член профсоюза:").grid(row=10, column=0, sticky=tk.W, padx=5, pady=2)
        self.teacher_union_member_var = tk.StringVar()
        ttk.Label(info_inner_frame, textvariable=self.teacher_union_member_var).grid(row=10, column=1, sticky=tk.W, padx=5, pady=2)
        
        # Кнопки быстрого доступа к функциям для выбранного преподавателя
        ttk.Separator(info_inner_frame, orient=tk.HORIZONTAL).grid(row=11, column=0, columnspan=2, sticky=tk.EW, pady=10)
        
        actions_frame = ttk.Frame(info_inner_frame)
        actions_frame.grid(row=12, column=0, columnspan=2, sticky=tk.EW, pady=5)
        
        ttk.Button(actions_frame, text="Рассчитать зарплату", command=self._calculate_salary_for_selected).pack(fill=tk.X, pady=2)
        ttk.Button(actions_frame, text="Запланировать отпуск", command=self._schedule_vacation_for_selected).pack(fill=tk.X, pady=2)

    
    def _get_teachers_list(self):
        """Возвращает список преподавателей в формате для комбобокса"""
        try:
            teachers = self.app.get_all_teachers()
            self.teachers_mapping = {t['name']: t['id'] for t in teachers}
            return [(t['id'], t['name']) for t in teachers]
        except Exception as e:
            logger.error(f"Ошибка получения списка преподавателей: {str(e)}")
            return []


    def _generate_report(self):
        """Единая функция для генерации отчетов разных типов"""
        report_type = self.report_type_var.get()
        
        try:
            if report_type == "salary":
                self._generate_salary_report()
            else:  # vacation
                self._generate_vacation_report()
        except Exception as e:
            logger.error(f"Ошибка при генерации отчета: {str(e)}", exc_info=True)
            messagebox.showerror("Ошибка", f"Ошибка при генерации отчета: {str(e)}")


    def _setup_salary_tab(self):
        """Настройка вкладки 'Расчет зарплаты'"""
        # Левая панель - выбор преподавателя и параметров расчета
        left_frame = ttk.Frame(self.salary_frame)
        left_frame = ttk.Frame(self.salary_frame, width=450)
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=False, padx=5, pady=5, ipadx=5, ipady=5)
        
        # Фрейм выбора преподавателя
        teacher_select_frame = ttk.LabelFrame(left_frame, text="Выбор преподавателя")
        teacher_select_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(teacher_select_frame, text="Преподаватель:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.salary_teacher_var = tk.StringVar()
        self.salary_teacher_combo = ttk.Combobox(teacher_select_frame, textvariable=self.salary_teacher_var, state="readonly", width=30)
        self.salary_teacher_combo.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)
        self.salary_teacher_combo.bind("<<ComboboxSelected>>", self._on_salary_teacher_select)
        
        # Фрейм параметров расчета
        calc_params_frame = ttk.LabelFrame(left_frame, text="Параметры расчета")
        calc_params_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Период расчета
        ttk.Label(calc_params_frame, text="Дата расчета:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.calc_date_var = tk.StringVar()
        self.calc_date_var.set(date.today().strftime("%d.%m.%Y"))
        calc_date_entry = ttk.Entry(calc_params_frame, textvariable=self.calc_date_var, width=15)
        calc_date_entry.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        ttk.Button(calc_params_frame, text="Календарь", command=lambda: self._show_calendar(self.calc_date_var)).grid(row=0, column=2, padx=5, pady=5)
        
        # Отработанные часы
        ttk.Label(calc_params_frame, text="Отработано часов:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.hours_worked_var = tk.StringVar()
        ttk.Entry(calc_params_frame, textvariable=self.hours_worked_var, width=10).grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Часы больничного
        ttk.Label(calc_params_frame, text="Часы больничного:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.sick_leave_hours_var = tk.StringVar()
        self.sick_leave_hours_var.set("0")
        ttk.Entry(calc_params_frame, textvariable=self.sick_leave_hours_var, width=10).grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Часы отсутствия
        ttk.Label(calc_params_frame, text="Часы отсутствия:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.absence_hours_var = tk.StringVar()
        self.absence_hours_var.set("0")
        ttk.Entry(calc_params_frame, textvariable=self.absence_hours_var, width=10).grid(row=3, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Бонусы
        ttk.Label(calc_params_frame, text="Бонус (руб):").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        self.bonus_var = tk.StringVar()
        self.bonus_var.set("0")
        ttk.Entry(calc_params_frame, textvariable=self.bonus_var, width=10).grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Ставка налога
        ttk.Label(calc_params_frame, text="Ставка налога (%):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
        self.tax_rate_var = tk.StringVar()
        self.tax_rate_var.set("13")
        ttk.Entry(calc_params_frame, textvariable=self.tax_rate_var, width=10).grid(row=5, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Кнопки расчета
        btn_frame = ttk.Frame(calc_params_frame)
        btn_frame.grid(row=6, column=0, columnspan=3, pady=10, sticky=tk.EW)
        
        ttk.Button(btn_frame, text="Рассчитать", command=self._calculate_salary).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Сохранить расчет", command=self._save_salary_calculation).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="Очистить", command=self._clear_salary_calculation).pack(side=tk.LEFT, padx=5)
        
        # Правая панель - результаты расчета
        right_frame = ttk.LabelFrame(self.salary_frame, text="Результаты расчета")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Таблица с результатами расчета
        results_frame = ttk.Frame(right_frame)
        results_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Основная информация о расчете
        ttk.Label(results_frame, text="Преподаватель:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_teacher_var = tk.StringVar()
        ttk.Label(results_frame, textvariable=self.result_teacher_var, font=("Arial", 10, "bold")).grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(results_frame, text="Дата расчета:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_date_var = tk.StringVar()
        ttk.Label(results_frame, textvariable=self.result_date_var).grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Separator(results_frame, orient=tk.HORIZONTAL).grid(row=2, column=0, columnspan=2, sticky=tk.EW, pady=5)
        
        # Детали расчета
        detail_frame = ttk.LabelFrame(results_frame, text="Детали расчета")
        detail_frame.grid(row=3, column=0, columnspan=2, sticky=tk.NSEW, pady=5)
        
        ttk.Label(detail_frame, text="Отработано часов:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_hours_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_hours_var).grid(row=0, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Часовая ставка:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_rate_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_rate_var).grid(row=1, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Базовая зарплата:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_base_salary_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_base_salary_var).grid(row=2, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Надбавка за должность:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_position_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_position_bonus_var).grid(row=3, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Надбавка за степень:").grid(row=4, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_degree_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_degree_bonus_var).grid(row=4, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Надбавка за стаж:").grid(row=5, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_experience_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_experience_bonus_var).grid(row=5, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Надбавка за категорию:").grid(row=6, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_category_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_category_bonus_var).grid(row=6, column=1, sticky=tk.W, padx=5, pady=2)
        

        ttk.Label(detail_frame, text="Надбавка молодому специалисту:").grid(row=7, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_young_specialist_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_young_specialist_bonus_var).grid(row=7, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Оплата больничных:").grid(row=8, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_sick_leave_pay_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_sick_leave_pay_var).grid(row=8, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Бонус:").grid(row=9, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_bonus_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_bonus_var).grid(row=9, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Separator(detail_frame, orient=tk.HORIZONTAL).grid(row=10, column=0, columnspan=2, sticky=tk.EW, pady=5)
        
        ttk.Label(detail_frame, text="Валовая зарплата:").grid(row=11, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_gross_salary_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_gross_salary_var, font=("Arial", 10, "bold")).grid(row=11, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Ставка налога:").grid(row=12, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_tax_rate_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_tax_rate_var).grid(row=12, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Сумма налога:").grid(row=13, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_tax_amount_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_tax_amount_var).grid(row=13, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Label(detail_frame, text="Профсоюзные взносы:").grid(row=14, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_union_contribution_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_union_contribution_var).grid(row=14, column=1, sticky=tk.W, padx=5, pady=2)
        
        ttk.Separator(detail_frame, orient=tk.HORIZONTAL).grid(row=15, column=0, columnspan=2, sticky=tk.EW, pady=5)
        
        ttk.Label(detail_frame, text="ЧИСТАЯ ЗАРПЛАТА:").grid(row=16, column=0, sticky=tk.W, padx=5, pady=2)
        self.result_net_salary_var = tk.StringVar()
        ttk.Label(detail_frame, textvariable=self.result_net_salary_var, font=("Arial", 12, "bold")).grid(row=16, column=1, sticky=tk.W, padx=5, pady=2)
        
        # Кнопки для экспорта результатов
        export_frame = ttk.Frame(results_frame)
        export_frame.grid(row=4, column=0, columnspan=2, sticky=tk.EW, pady=10)
        
        ttk.Button(export_frame, text="Экспорт в PDF", command=self._export_salary_to_pdf).pack(side=tk.LEFT, padx=5)
        ttk.Button(export_frame, text="Печать", command=self._print_salary_result).pack(side=tk.LEFT, padx=5)
    
    # Добавьте этот метод после метода _setup_salary_tab или в любом удобном месте класса
    def _show_calendar(self, string_var):
        """Отображает календарь для выбора даты и записывает результат в переменную StringVar"""
        top = tk.Toplevel(self.master)
        top.title("Выберите дату")
        
        def set_date(event):
            date_str = cal.get_date()
            string_var.set(date_str)
            top.destroy()
        
        cal = tkcalendar.Calendar(top, selectmode='day', 
                                date_pattern='dd.MM.yyyy',
                                background='darkblue', 
                                foreground='white',
                                borderwidth=2)
        cal.pack(padx=10, pady=10)
        cal.bind("<<CalendarSelected>>", set_date)
        
        ttk.Button(top, text="Выбрать", command=lambda: set_date(None)).pack(pady=5)

    def get_vacation_tab(self):
        """Возвращает объект управления вкладкой отпусков"""
        return self.vacation_tab_manager if hasattr(self, 'vacation_tab_manager') else None
    
    def _setup_vacation_tab(self):
        """Настройка вкладки 'Отпуска' с использованием класса VacationTab"""
        try:
            # Используем существующее подключение к БД из self.app
            db_conn = self.app.db_connection
            
            # Инициализируем менеджер вкладки отпусков
            self.vacation_tab_manager = VacationTab(self.notebook, db_conn)
            
            # Удаляем созданную ранее вкладку, так как VacationTab создаст свою собственную
            self.notebook.forget(self.vacation_frame)
            
            logger.info("Вкладка отпусков успешно инициализирована")
        except Exception as e:
            logger.error(f"Ошибка при инициализации вкладки отпусков: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось инициализировать вкладку отпусков: {str(e)}")
            # Создаем простую вкладку с сообщением об ошибке
            error_frame = ttk.Frame(self.notebook)
            self.notebook.add(error_frame, text="Отпуска")
            ttk.Label(error_frame, text=f"Ошибка загрузки: {str(e)}", foreground="red").pack(padx=20, pady=20)

    def _init_report_combobox(self):
        """Инициализация комбобокса отчетов"""
        if self.report_teacher_combo is None:
            self.report_teacher_combo = ttk.Combobox(self.reports_frame)
            self.report_teacher_combo.grid(row=0, column=1, padx=5, pady=5)
        
        teachers = self._get_teachers_list()
        self.report_teacher_combo['values'] = [name for _, name in teachers]

    def _setup_reports_tab(self):
        """Настройка вкладки 'Отчеты'"""
        # Основной фрейм
        main_frame = ttk.Frame(self.reports_frame)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)
        
        # Фрейм выбора типа отчета
        report_type_frame = ttk.LabelFrame(main_frame, text="Тип отчета")
        report_type_frame.pack(fill=tk.X, padx=5, pady=5)
        
        self.report_type_var = tk.StringVar(value="salary")
        
        ttk.Radiobutton(
            report_type_frame,
            text="Отчет по зарплате",
            variable=self.report_type_var,
            value="salary",
            command=self._on_report_type_change
        ).grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        
        ttk.Radiobutton(
            report_type_frame,
            text="Отчет по отпускам",
            variable=self.report_type_var,
            value="vacation",
            command=self._on_report_type_change
        ).grid(row=0, column=1, sticky=tk.W, padx=10, pady=5)
        
        # Фрейм для параметров отчета
        self.report_params_frame = ttk.LabelFrame(main_frame, text="Параметры отчета")
        self.report_params_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Параметры отчета по зарплате
        self.salary_report_frame = ttk.Frame(self.report_params_frame)
        self.salary_report_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        
        # Выбор преподавателя для отчета
        teacher_frame = ttk.Frame(self.salary_report_frame)
        teacher_frame.pack(fill=tk.X, padx=5, pady=5)
        
        # Переключатель между индивидуальным и общим отчетом
        self.salary_report_scope_var = tk.StringVar(value="individual")
        ttk.Radiobutton(
            teacher_frame,
            text="Индивидуальный отчет",
            variable=self.salary_report_scope_var,
            value="individual",
            command=self._update_salary_report_options
        ).grid(row=0, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        ttk.Radiobutton(
            teacher_frame,
            text="Общий отчет по всем преподавателям",
            variable=self.salary_report_scope_var,
            value="all",
            command=self._update_salary_report_options
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        # Комбобокс для выбора преподавателя
        ttk.Label(teacher_frame, text="Преподаватель:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        
        # Получаем список преподавателей
        teachers = self._get_teachers_list()
        teacher_names = [name for _, name in teachers]
        
        self.report_teacher_var = tk.StringVar()
        self.report_teacher_combo = ttk.Combobox(
            teacher_frame,
            textvariable=self.report_teacher_var,
            values=teacher_names,
            state="readonly",
            width=40
        )
        self.report_teacher_combo.grid(row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Если есть текущий выбранный преподаватель, используем его
        if self.current_teacher_id is not None:
            for teacher_id, name in teachers:
                if teacher_id == self.current_teacher_id:
                    self.report_teacher_var.set(name)
                    break
        
        # Фрейм для выбора периода отчета
        period_frame = ttk.Frame(self.salary_report_frame)
        period_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(period_frame, text="Начало периода:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.start_date_var = tk.StringVar(value=date.today().replace(day=1).strftime("%d.%m.%Y"))
        start_date_entry = ttk.Entry(period_frame, textvariable=self.start_date_var, width=15)
        start_date_entry.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        ttk.Button(
            period_frame,
            text="...",
            command=lambda: self._show_calendar(self.start_date_var),
            width=3
        ).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Label(period_frame, text="Конец периода:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.end_date_var = tk.StringVar(value=date.today().strftime("%d.%m.%Y"))
        end_date_entry = ttk.Entry(period_frame, textvariable=self.end_date_var, width=15)
        end_date_entry.grid(row=1, column=1, sticky=tk.W, padx=5, pady=5)
        ttk.Button(
            period_frame,
            text="...",
            command=lambda: self._show_calendar(self.end_date_var),
            width=3
        ).grid(row=1, column=2, padx=5, pady=5)
        
        # Формат и опции отчета
        options_frame = ttk.Frame(self.salary_report_frame)
        options_frame.pack(fill=tk.X, padx=5, pady=5)
        
        ttk.Label(options_frame, text="Формат отчета:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.report_format_var = tk.StringVar(value="pdf")
        format_combo = ttk.Combobox(
            options_frame,
            textvariable=self.report_format_var,
            values=["pdf", "excel", "csv"],
            state="readonly",
            width=15
        )
        format_combo.grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Дополнительные опции
        self.include_details_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(
            options_frame,
            text="Включить детали расчета",
            variable=self.include_details_var
        ).grid(row=1, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        
        # Параметры отчета по отпускам (скрыты по умолчанию)
        self.vacation_report_frame = ttk.Frame(self.report_params_frame)
        # По умолчанию отображаем отчет по зарплате
        
        # Фрейм для кнопок
        buttons_frame = ttk.Frame(main_frame)
        buttons_frame.pack(fill=tk.X, padx=5, pady=10)
        
        ttk.Button(
            buttons_frame,
            text="Сформировать отчет",
            command=self._generate_report
        ).pack(side=tk.RIGHT, padx=5)
        
        # Инициализируем состояние интерфейса
        self._update_salary_report_options()

    def _on_report_type_change(self):
        """Обработчик изменения типа отчета"""
        report_type = self.report_type_var.get()
        
        if report_type == "salary":
            self.vacation_report_frame.pack_forget()
            self.salary_report_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        else:  # vacation
            self.salary_report_frame.pack_forget()
            self.vacation_report_frame.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
            # Здесь можно добавить код для настройки параметров отчета по отпускам

    def _update_salary_report_options(self):
        """Обновляет доступность опций в зависимости от выбранного типа отчета по зарплате"""
        scope = self.salary_report_scope_var.get()
        
        if scope == "individual":
            self.report_teacher_combo.config(state="readonly")
        else:  # all
            self.report_teacher_combo.config(state="disabled")

    
    def update_status(self, message):
        """Обновить статусную строку"""
        if hasattr(self, 'status_var'):
            self.status_var.set(message)
        else:
            print(f"Статус: {message}")  # Запасной вариант, если status_var не доступен

    def _on_close(self):
        """Обработчик события закрытия окна"""
        try:
            # Корректно закрываем соединения с базой данных
            self.app.close()
            self.master.destroy()
        except Exception as e:
            logger.error(f"Ошибка при закрытии приложения: {str(e)}")
            self.master.destroy()
    
    def _show_calendar(self, date_var):
        """Показать диалог выбора даты"""
        def set_date():
            date_var.set(cal.get_date().strftime("%d.%m.%Y"))
            top.destroy()
        
        top = tk.Toplevel(self.master)
        top.title("Выберите дату")
        top.grab_set()  # Модальное окно
        
        # Если в поле уже есть дата, используем её
        try:
            current_date = datetime.datetime.strptime(date_var.get(), "%d.%m.%Y").date()
        except:
            current_date = date.today()
        
        cal = tkcalendar.Calendar(top, selectmode='day', year=current_date.year, 
                                month=current_date.month, day=current_date.day, 
                                locale='ru_RU', date_pattern="dd.mm.yyyy")
        cal.pack(padx=10, pady=10)
        
        ttk.Button(top, text="Выбрать", command=set_date).pack(pady=5)
    
    def _show_db_settings(self):
        """Показать диалог настроек подключения к базе данных"""
        settings_window = tk.Toplevel(self.master)
        settings_window.title("Настройки базы данных")
        settings_window.geometry("400x250")
        settings_window.resizable(False, False)
        settings_window.grab_set()
        
        ttk.Label(settings_window, text="Хост:").grid(row=0, column=0, sticky=tk.W, padx=10, pady=5)
        host_var = tk.StringVar(value=self.db_config['host'])
        ttk.Entry(settings_window, textvariable=host_var, width=30).grid(row=0, column=1, padx=10, pady=5)
        
        ttk.Label(settings_window, text="База данных:").grid(row=1, column=0, sticky=tk.W, padx=10, pady=5)
        db_var = tk.StringVar(value=self.db_config['database'])
        ttk.Entry(settings_window, textvariable=db_var, width=30).grid(row=1, column=1, padx=10, pady=5)
        
        ttk.Label(settings_window, text="Пользователь:").grid(row=2, column=0, sticky=tk.W, padx=10, pady=5)
        user_var = tk.StringVar(value=self.db_config['user'])
        ttk.Entry(settings_window, textvariable=user_var, width=30).grid(row=2, column=1, padx=10, pady=5)
        
        ttk.Label(settings_window, text="Пароль:").grid(row=3, column=0, sticky=tk.W, padx=10, pady=5)
        password_var = tk.StringVar(value=self.db_config['password'])
        password_entry = ttk.Entry(settings_window, textvariable=password_var, width=30, show="*")
        password_entry.grid(row=3, column=1, padx=10, pady=5)
        
        ttk.Label(settings_window, text="Порт:").grid(row=4, column=0, sticky=tk.W, padx=10, pady=5)
        port_var = tk.StringVar(value=self.db_config['port'])
        ttk.Entry(settings_window, textvariable=port_var, width=10).grid(row=4, column=1, sticky=tk.W, padx=10, pady=5)
        
        def save_settings():
            try:
                # Обновляем конфигурацию
                self.db_config['host'] = host_var.get()
                self.db_config['database'] = db_var.get()
                self.db_config['user'] = user_var.get()
                self.db_config['password'] = password_var.get()
                self.db_config['port'] = port_var.get()
                
                # Пересоздаем соединение с базой данных
                self.app.close()
                self.app = SalaryApp(self.db_config)
                
                # Обновляем данные
                self._load_teachers()
                
                messagebox.showinfo("Успех", "Настройки успешно сохранены")
                settings_window.destroy()
            except Exception as e:
                logger.error(f"Ошибка при сохранении настроек БД: {str(e)}")
                messagebox.showerror("Ошибка", f"Не удалось сохранить настройки: {str(e)}")
        
        def test_connection():
            try:
                import psycopg2
                conn = psycopg2.connect(
                    host=host_var.get(),
                    database=db_var.get(),
                    user=user_var.get(),
                    password=password_var.get(),
                    port=port_var.get()
                )
                conn.close()
                messagebox.showinfo("Успех", "Соединение с базой данных успешно установлено")
            except Exception as e:
                logger.error(f"Ошибка при тестовом подключении к БД: {str(e)}")
                messagebox.showerror("Ошибка", f"Не удалось подключиться к базе данных: {str(e)}")
        
        btn_frame = ttk.Frame(settings_window)
        btn_frame.grid(row=5, column=0, columnspan=2, pady=20)
        
        ttk.Button(btn_frame, text="Проверить соединение", command=test_connection).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Сохранить", command=save_settings).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Отмена", command=settings_window.destroy).pack(side=tk.LEFT, padx=10)
    
    def _show_about(self):
        """Показать информацию о программе"""
        about_text = """
        Система расчета зарплаты для лицея ПолессГУ
        Версия: 1.0.0
        
        Данная программа предназначена для автоматизации расчета заработной платы, 
        отпусков и больничных в системе образования Республики Беларусь.
        
        © 2025 Все права защищены
        """
        
        about_window = tk.Toplevel(self.master)
        about_window.title("О программе")
        about_window.geometry("500x300")
        about_window.resizable(False, False)
        about_window.grab_set()
        
        ttk.Label(about_window, text="Система расчета зарплаты", font=("Arial", 16, "bold")).pack(pady=10)
        ttk.Label(about_window, text="для лицея ПолессГУ", font=("Arial", 14)).pack()
        
        text_widget = tk.Text(about_window, wrap=tk.WORD, height=10, width=50)
        text_widget.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
        text_widget.insert(tk.END, about_text)
        text_widget.config(state=tk.DISABLED)
        
        ttk.Button(about_window, text="OK", command=about_window.destroy).pack(pady=10)
    
    def _show_user_manual(self):
        """Показать руководство пользователя"""
        manual_text = """
        РУКОВОДСТВО ПОЛЬЗОВАТЕЛЯ
        
        1. Вкладка "Преподаватели"
           - Здесь вы можете управлять списком преподавателей
           - Добавлять, редактировать и удалять записи
           - Просматривать детальную информацию
        
        2. Вкладка "Расчет зарплаты"
           - Выберите преподавателя и укажите параметры расчета
           - Нажмите "Рассчитать" для просмотра результата
           - Сохраните расчет в базу данных
        
        3. Вкладка "Отпуска"
           - Управление отпусками преподавателей
           - Планирование и расчет отпускных выплат
           - Оптимальное распределение отпуска
        
        
        4. Вкладка "Отчеты"
           - Формирование различных типов отчетов
           - Экспорт отчетов в файл
        """
        
        manual_window = tk.Toplevel(self.master)
        manual_window.title("Руководство пользователя")
        manual_window.geometry("700x500")
        manual_window.grab_set()
        
        text_widget = tk.Text(manual_window, wrap=tk.WORD)
        text_widget.pack(padx=20, pady=20, fill=tk.BOTH, expand=True)
        
        scrollbar = ttk.Scrollbar(text_widget, orient=tk.VERTICAL, command=text_widget.yview)
        text_widget.configure(yscroll=scrollbar.set)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        
        text_widget.insert(tk.END, manual_text)
        text_widget.config(state=tk.DISABLED)
        
        ttk.Button(manual_window, text="Закрыть", command=manual_window.destroy).pack(pady=10)
    
    # --- Методы для работы с преподавателями ---
    
# Замените или добавьте этот метод _load_teachers, с проверками на None:
    def _load_teachers(self):
        """Загрузка списка преподавателей в таблицу"""
        # Очистить таблицу
        for item in self.teachers_tree.get_children():
            self.teachers_tree.delete(item)
        
        try:
            # Получить список преподавателей
            teachers = self.app.get_all_teachers()
            
            if not teachers:
                self.update_status("Список преподавателей пуст")
                return
            
            # Заполнить таблицу
            for teacher in teachers:
                if teacher is None:
                    continue  # Пропускаем None значения
                    
                # Безопасное получение значений из словаря
                teacher_id = teacher.get('id', '')
                name = teacher.get('name', '')
                position = teacher.get('position', '')
                degree = teacher.get('academic_degree', '')
                category = teacher.get('qualification_category', '')
                rate = teacher.get('hourly_rate', '')
                experience = teacher.get('experience_years', '')
                
                self.teachers_tree.insert('', 'end', values=(
                    teacher_id, name, position, degree, category, rate, experience
                ))
            
            # Обновить комбобоксы
            teachers_list = [(t.get('id', ''), t.get('name', '')) for t in teachers if t is not None]
            
            # Безопасное обновление комбобоксов
            if hasattr(self, 'salary_teacher_combo'):
                self.salary_teacher_combo['values'] = [name for _, name in teachers_list]
        
            
            if hasattr(self, 'report_teacher_combo') and self.report_teacher_combo:
                self.report_teacher_combo['values'] = [name for _, name in teachers_list]
            
            self.update_status(f"Загружено {len(teachers)} преподавателей")
        except Exception as e:
            logger.error(f"Ошибка при загрузке преподавателей: {str(e)}")
            self.update_status(f"Ошибка: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при загрузке преподавателей: {str(e)}")

    def _update_teacher_comboboxes(self, teachers=None):
        """Обновить списки преподавателей в комбобоксах"""
        if teachers is None:
            try:
                teachers = self.app.get_all_teachers()
            except Exception as e:
                logger.error(f"Ошибка при получении преподавателей: {str(e)}")
                return
        
        # Создаем список для комбобоксов
        teacher_list = [f"{t['id']} - {t['name']}" for t in teachers]
        
        # Обновляем все комбобоксы
        self.salary_teacher_combo['values'] = teacher_list
        self.report_teacher_combo['values'] = teacher_list
    
    def _on_teacher_select(self, event):
        """Обработчик выбора преподавателя в таблице"""
        selected_items = self.teachers_tree.selection()
        if not selected_items:
            return
        
        item = selected_items[0]
        teacher_id = self.teachers_tree.item(item, 'values')[0]
        
        try:
            # Получаем данные преподавателя
            teacher = self.app.get_teacher_by_id(int(teacher_id))
            if not teacher:
                return
            
            # Сохраняем ID текущего выбранного преподавателя
            self.current_teacher_id = teacher['id']
            
            # Обновляем информацию в правой панели
            self.teacher_id_var.set(str(teacher['id']))
            self.teacher_name_var.set(teacher['name'])
            self.teacher_position_var.set(teacher.get('position', 'Не указана'))
            self.teacher_degree_var.set(teacher.get('academic_degree', 'Нет'))
            self.teacher_qualification_var.set(teacher.get('qualification_category', 'Нет'))
            self.teacher_experience_var.set(str(teacher.get('experience_years', '0')))
            self.teacher_rate_var.set(str(teacher.get('hourly_rate', '0')))
            
            # Форматирование дат
            hire_date = teacher.get('hire_date')
            if hire_date:
                if isinstance(hire_date, str):
                    self.teacher_hire_date_var.set(hire_date)
                else:
                    self.teacher_hire_date_var.set(hire_date.strftime("%d.%m.%Y"))
            else:
                self.teacher_hire_date_var.set("Не указана")
            
            birth_date = teacher.get('birth_date')
            if birth_date:
                if isinstance(birth_date, str):
                    self.teacher_birth_date_var.set(birth_date)
                else:
                    self.teacher_birth_date_var.set(birth_date.strftime("%d.%m.%Y"))
            else:
                self.teacher_birth_date_var.set("Не указана")
            
            # Статусы "да/нет"
            self.teacher_young_specialist_var.set("Да" if teacher.get('is_young_specialist', False) else "Нет")
            self.teacher_union_member_var.set("Да" if teacher.get('is_union_member', False) else "Нет")
            
            self.update_status(f"Выбран преподаватель: {teacher['name']}")
        except Exception as e:
            logger.error(f"Ошибка при загрузке данных преподавателя: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось загрузить данные преподавателя: {str(e)}")
    
    def _show_add_teacher_dialog(self):
        """Показать диалог добавления преподавателя"""
        dialog = tk.Toplevel(self.master)
        dialog.title("Добавление преподавателя")
        dialog.geometry("600x500")
        dialog.grab_set()
        
        # Создаем форму для ввода данных
        form_frame = ttk.Frame(dialog, padding=10)
        form_frame.pack(fill=tk.BOTH, expand=True)
        
        # ФИО
        ttk.Label(form_frame, text="ФИО:*").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        name_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=name_var, width=40).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Должность
        ttk.Label(form_frame, text="Должность:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        position_var = tk.StringVar()
        positions = ["ассистент", "преподаватель", "старший преподаватель", "доцент", "профессор", 
                    "заведующий кафедрой", "декан"]
        ttk.Combobox(form_frame, textvariable=position_var, values=positions, width=30).grid(
            row=1, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Ученая степень
        ttk.Label(form_frame, text="Ученая степень:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        degree_var = tk.StringVar()
        degrees = ["", "кандидат наук", "доктор наук"]
        ttk.Combobox(form_frame, textvariable=degree_var, values=degrees, width=20).grid(
            row=2, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Квалификационная категория
        ttk.Label(form_frame, text="Квалификация:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        qualification_var = tk.StringVar()
        qualifications = ["", "вторая", "первая", "высшая"]
        ttk.Combobox(form_frame, textvariable=qualification_var, values=qualifications, width=15).grid(
            row=3, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Ставка
        ttk.Label(form_frame, text="Часовая ставка:*").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        rate_var = tk.StringVar()
        ttk.Entry(form_frame, textvariable=rate_var, width=10).grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Стаж
        ttk.Label(form_frame, text="Стаж (лет):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
        experience_var = tk.StringVar()
        experience_var.set("0")
        ttk.Entry(form_frame, textvariable=experience_var, width=5).grid(row=5, column=1, sticky=tk.W, padx=5, pady=5)
        
        # Дата приема на работу
        ttk.Label(form_frame, text="Дата приема:*").grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
        hire_date_var = tk.StringVar()
        hire_date_var.set(date.today().strftime("%d.%m.%Y"))
        hire_date_frame = ttk.Frame(form_frame)
        hire_date_frame.grid(row=6, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Entry(hire_date_frame, textvariable=hire_date_var, width=15).pack(side=tk.LEFT)
        ttk.Button(hire_date_frame, text="Календарь", 
                  command=lambda: self._show_calendar(hire_date_var)).pack(side=tk.LEFT, padx=5)
        
        # Дата рождения
        ttk.Label(form_frame, text="Дата рождения:").grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
        birth_date_var = tk.StringVar()
        birth_date_frame = ttk.Frame(form_frame)
        birth_date_frame.grid(row=7, column=1, sticky=tk.W, padx=5, pady=5)
        
        ttk.Entry(birth_date_frame, textvariable=birth_date_var, width=15).pack(side=tk.LEFT)
        ttk.Button(birth_date_frame, text="Календарь", 
                  command=lambda: self._show_calendar(birth_date_var)).pack(side=tk.LEFT, padx=5)
        
        # Молодой специалист
        young_specialist_var = tk.BooleanVar()
        ttk.Checkbutton(form_frame, text="Молодой специалист", 
                       variable=young_specialist_var).grid(row=8, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        # Член профсоюза
        union_member_var = tk.BooleanVar()
        ttk.Checkbutton(form_frame, text="Член профсоюза", 
                       variable=union_member_var).grid(row=9, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        ttk.Label(form_frame, text="* - обязательные поля").grid(row=10, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
        
        # Кнопки
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(fill=tk.X, padx=10, pady=10)
        
        def add_teacher():
            # Проверка обязательных полей
            if not name_var.get().strip():
                messagebox.showerror("Ошибка", "Поле 'ФИО' обязательно для заполнения")
                return
            
            if not rate_var.get().strip():
                messagebox.showerror("Ошибка", "Поле 'Часовая ставка' обязательно для заполнения")
                return
            
            if not hire_date_var.get().strip():
                messagebox.showerror("Ошибка", "Поле 'Дата приема' обязательно для заполнения")
                return
            
            try:
                # Преобразование даты
                hire_date = datetime.datetime.strptime(hire_date_var.get(), "%d.%m.%Y").date()
                
                birth_date = None
                if birth_date_var.get().strip():
                    birth_date = datetime.datetime.strptime(birth_date_var.get(), "%d.%m.%Y").date()
                
                # Создаем словарь с данными преподавателя
                teacher_data = {
                    'name': name_var.get().strip(),
                    'position': position_var.get(),
                    'academic_degree': degree_var.get(),
                    'qualification_category': qualification_var.get(),
                    'hourly_rate': float(rate_var.get()),
                    'experience_years': int(experience_var.get()),
                    'hire_date': hire_date,
                    'birth_date': birth_date,
                    'is_young_specialist': young_specialist_var.get(),
                    'is_union_member': union_member_var.get()
                }
                
                # Добавляем преподавателя
                teacher_id = self.app.add_teacher(teacher_data)
                
                # Обновляем список преподавателей
                self._load_teachers()
                
                messagebox.showinfo("Успех", f"Преподаватель успешно добавлен (ID: {teacher_id})")
                dialog.destroy()
            except Exception as e:
                logger.error(f"Ошибка при добавлении преподавателя: {str(e)}")
                messagebox.showerror("Ошибка", f"Не удалось добавить преподавателя: {str(e)}")
        
        ttk.Button(btn_frame, text="Добавить", command=add_teacher).pack(side=tk.LEFT, padx=10)
        ttk.Button(btn_frame, text="Отмена", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)
    
    def _show_edit_teacher_dialog(self):
        """Показать диалог редактирования преподавателя"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель для редактирования")
            return
        
        try:
            # Получаем данные преподавателя
            teacher = self.app.get_teacher_by_id(self.current_teacher_id)
            if not teacher:
                messagebox.showerror("Ошибка", "Не удалось загрузить данные преподавателя")
                return
            
            dialog = tk.Toplevel(self.master)
            dialog.title(f"Редактирование преподавателя: {teacher['name']}")
            dialog.geometry("600x500")
            dialog.grab_set()
            
            # Создаем форму для ввода данных
            form_frame = ttk.Frame(dialog, padding=10)
            form_frame.pack(fill=tk.BOTH, expand=True)
            
            # ФИО
            ttk.Label(form_frame, text="ФИО:*").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
            name_var = tk.StringVar(value=teacher['name'])
            ttk.Entry(form_frame, textvariable=name_var, width=40).grid(row=0, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Должность
            ttk.Label(form_frame, text="Должность:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
            position_var = tk.StringVar(value=teacher.get('position', ''))

            positions = ["ассистент", "преподаватель", "старший преподаватель", "доцент", "профессор", 
                        "заведующий кафедрой", "декан"]
            ttk.Combobox(form_frame, textvariable=position_var, values=positions, width=30).grid(
                row=1, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Ученая степень
            ttk.Label(form_frame, text="Ученая степень:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
            degree_var = tk.StringVar(value=teacher.get('academic_degree', ''))
            degrees = ["", "кандидат наук", "доктор наук"]
            ttk.Combobox(form_frame, textvariable=degree_var, values=degrees, width=20).grid(
                row=2, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Квалификационная категория
            ttk.Label(form_frame, text="Квалификация:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
            qualification_var = tk.StringVar(value=teacher.get('qualification_category', ''))
            qualifications = ["", "вторая", "первая", "высшая"]
            ttk.Combobox(form_frame, textvariable=qualification_var, values=qualifications, width=15).grid(
                row=3, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Ставка
            ttk.Label(form_frame, text="Часовая ставка:*").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
            rate_var = tk.StringVar(value=str(teacher.get('hourly_rate', '')))
            ttk.Entry(form_frame, textvariable=rate_var, width=10).grid(row=4, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Стаж
            ttk.Label(form_frame, text="Стаж (лет):").grid(row=5, column=0, sticky=tk.W, padx=5, pady=5)
            experience_var = tk.StringVar(value=str(teacher.get('experience_years', 0)))
            ttk.Entry(form_frame, textvariable=experience_var, width=5).grid(row=5, column=1, sticky=tk.W, padx=5, pady=5)
            
            # Дата приема на работу
            ttk.Label(form_frame, text="Дата приема:*").grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
            hire_date_var = tk.StringVar()
            if teacher.get('hire_date'):
                if isinstance(teacher['hire_date'], str):
                    hire_date_var.set(teacher['hire_date'])
                else:
                    hire_date_var.set(teacher['hire_date'].strftime("%d.%m.%Y"))
            
            hire_date_frame = ttk.Frame(form_frame)
            hire_date_frame.grid(row=6, column=1, sticky=tk.W, padx=5, pady=5)
            
            ttk.Entry(hire_date_frame, textvariable=hire_date_var, width=15).pack(side=tk.LEFT)
            ttk.Button(hire_date_frame, text="Календарь", 
                      command=lambda: self._show_calendar(hire_date_var)).pack(side=tk.LEFT, padx=5)
            
            # Дата рождения
            ttk.Label(form_frame, text="Дата рождения:").grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
            birth_date_var = tk.StringVar()
            if teacher.get('birth_date'):
                if isinstance(teacher['birth_date'], str):
                    birth_date_var.set(teacher['birth_date'])
                else:
                    birth_date_var.set(teacher['birth_date'].strftime("%d.%m.%Y"))
            
            birth_date_frame = ttk.Frame(form_frame)
            birth_date_frame.grid(row=7, column=1, sticky=tk.W, padx=5, pady=5)
            
            ttk.Entry(birth_date_frame, textvariable=birth_date_var, width=15).pack(side=tk.LEFT)
            ttk.Button(birth_date_frame, text="Календарь", 
                      command=lambda: self._show_calendar(birth_date_var)).pack(side=tk.LEFT, padx=5)
            
            # Молодой специалист
            young_specialist_var = tk.BooleanVar(value=teacher.get('is_young_specialist', False))
            ttk.Checkbutton(form_frame, text="Молодой специалист", 
                           variable=young_specialist_var).grid(row=8, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
            
            # Член профсоюза
            union_member_var = tk.BooleanVar(value=teacher.get('is_union_member', False))
            ttk.Checkbutton(form_frame, text="Член профсоюза", 
                           variable=union_member_var).grid(row=9, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
            
            ttk.Label(form_frame, text="* - обязательные поля").grid(row=10, column=0, columnspan=2, sticky=tk.W, padx=5, pady=5)
            
            # Кнопки
            btn_frame = ttk.Frame(dialog)
            btn_frame.pack(fill=tk.X, padx=10, pady=10)
            
            def update_teacher():
                # Проверка обязательных полей
                if not name_var.get().strip():
                    messagebox.showerror("Ошибка", "Поле 'ФИО' обязательно для заполнения")
                    return
                
                if not rate_var.get().strip():
                    messagebox.showerror("Ошибка", "Поле 'Часовая ставка' обязательно для заполнения")
                    return
                
                if not hire_date_var.get().strip():
                    messagebox.showerror("Ошибка", "Поле 'Дата приема' обязательно для заполнения")
                    return
                
                try:
                    # Преобразование даты
                    hire_date = datetime.datetime.strptime(hire_date_var.get(), "%d.%m.%Y").date()
                    
                    birth_date = None
                    if birth_date_var.get().strip():
                        birth_date = datetime.datetime.strptime(birth_date_var.get(), "%d.%m.%Y").date()
                    
                    # Создаем словарь с данными преподавателя
                    teacher_data = {
                        'name': name_var.get().strip(),
                        'position': position_var.get(),
                        'academic_degree': degree_var.get(),
                        'qualification_category': qualification_var.get(),
                        'hourly_rate': float(rate_var.get()),
                        'experience_years': int(experience_var.get()),
                        'hire_date': hire_date,
                        'birth_date': birth_date,
                        'is_young_specialist': young_specialist_var.get(),
                        'is_union_member': union_member_var.get()
                    }
                    
                    # Обновляем данные преподавателя
                    success = self.app.update_teacher(self.current_teacher_id, teacher_data)
                    
                    if success:
                        # Обновляем список преподавателей
                        self._load_teachers()
                        
                        # Обновляем информацию в правой панели
                        self._on_teacher_select(None)
                        
                        messagebox.showinfo("Успех", "Данные преподавателя успешно обновлены")
                        dialog.destroy()
                    else:
                        messagebox.showerror("Ошибка", "Не удалось обновить данные преподавателя")
                except Exception as e:
                    logger.error(f"Ошибка при обновлении данных преподавателя: {str(e)}")
                    messagebox.showerror("Ошибка", f"Не удалось обновить данные преподавателя: {str(e)}")
            
            ttk.Button(btn_frame, text="Сохранить", command=update_teacher).pack(side=tk.LEFT, padx=10)
            ttk.Button(btn_frame, text="Отмена", command=dialog.destroy).pack(side=tk.RIGHT, padx=10)
        except Exception as e:
            logger.error(f"Ошибка при открытии диалога редактирования: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось открыть диалог редактирования: {str(e)}")
    
    def _delete_teacher(self):
        """Удалить выбранного преподавателя"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель для удаления")
            return
        
        try:
            # Получаем данные преподавателя
            teacher = self.app.get_teacher_by_id(self.current_teacher_id)
            if not teacher:
                messagebox.showerror("Ошибка", "Не удалось загрузить данные преподавателя")
                return
            
            # Запрашиваем подтверждение
            confirm = messagebox.askyesno(
                "Подтверждение удаления", 
                f"Вы действительно хотите удалить преподавателя {teacher['name']}?\n\n" +
                "Внимание: Это действие приведет к удалению всех связанных с ним данных " +
                "(расчеты зарплаты, отпуска)."
            )
            
            if not confirm:
                return
            
            # Удаляем преподавателя
            success = self.app.delete_teacher(self.current_teacher_id)
            
            if success:
                # Обновляем список преподавателей
                self._load_teachers()
                
                # Очищаем информацию в правой панели
                self.current_teacher_id = None
                self.teacher_id_var.set("")
                self.teacher_name_var.set("")
                self.teacher_position_var.set("")
                self.teacher_degree_var.set("")
                self.teacher_qualification_var.set("")
                self.teacher_experience_var.set("")
                self.teacher_rate_var.set("")
                self.teacher_hire_date_var.set("")
                self.teacher_birth_date_var.set("")
                self.teacher_young_specialist_var.set("")
                self.teacher_union_member_var.set("")
                
                messagebox.showinfo("Успех", f"Преподаватель {teacher['name']} успешно удален")
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить преподавателя")
        except Exception as e:
            logger.error(f"Ошибка при удалении преподавателя: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось удалить преподавателя: {str(e)}")
    
    def _on_search(self, event):
        """Обработчик поиска преподавателей"""
        search_term = self.search_var.get().lower()
        
        # Очищаем текущий список
        for item in self.teachers_tree.get_children():
            self.teachers_tree.delete(item)
        
        try:
            # Получаем список всех преподавателей
            all_teachers = self.app.get_all_teachers()
            
            # Фильтруем список по поисковому запросу
            if search_term:
                filtered_teachers = [t for t in all_teachers if search_term in t['name'].lower() or 
                                  (t.get('position') and search_term in t['position'].lower())]
            else:
                filtered_teachers = all_teachers
            
            # Заполняем таблицу
            for teacher in filtered_teachers:
                self.teachers_tree.insert('', tk.END, values=(
                    teacher['id'],
                    teacher['name'],
                    teacher.get('position', ''),
                    teacher.get('academic_degree', ''),
                    teacher.get('qualification_category', ''),
                    teacher.get('hourly_rate', ''),
                    teacher.get('experience_years', '')
                ))
            
            self.update_status(f"Найдено преподавателей: {len(filtered_teachers)}")
        except Exception as e:
            logger.error(f"Ошибка при поиске преподавателей: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при поиске: {str(e)}")
    
    def _import_teachers_from_csv(self):
        """Импорт преподавателей из CSV-файла"""
        file_path = filedialog.askopenfilename(
            title="Выберите CSV-файл для импорта",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            import csv
            
            teachers_added = 0
            teachers_updated = 0
            errors = []
            
            with open(file_path, 'r', encoding='utf-8') as csvfile:
                reader = csv.DictReader(csvfile)
                
                for i, row in enumerate(reader, start=2):  # start=2 because row 1 is headers
                    try:
                        # Проверка обязательных полей
                        if not row.get('name') or not row.get('hourly_rate') or not row.get('hire_date'):
                            errors.append(f"Строка {i}: Отсутствуют обязательные поля")
                            continue
                        
                        # Преобразование типов данных
                        teacher_data = {
                            'name': row['name'].strip(),
                            'position': row.get('position', '').strip(),
                            'academic_degree': row.get('academic_degree', '').strip(),
                            'qualification_category': row.get('qualification_category', '').strip(),
                            'hourly_rate': float(row['hourly_rate']),
                            'experience_years': int(row.get('experience_years', 0)),
                            'hire_date': datetime.datetime.strptime(row['hire_date'], "%d.%m.%Y").date(),
                            'is_young_specialist': row.get('is_young_specialist', '').lower() in ('true', 'yes', '1', 'да'),
                            'is_union_member': row.get('is_union_member', '').lower() in ('true', 'yes', '1', 'да')
                        }
                        
                        if row.get('birth_date'):
                            teacher_data['birth_date'] = datetime.datetime.strptime(row['birth_date'], "%d.%m.%Y").date()
                        
                        # Если указан ID - обновляем, иначе добавляем
                        if row.get('id'):
                            success = self.app.update_teacher(int(row['id']), teacher_data)
                            if success:
                                teachers_updated += 1
                        else:
                            self.app.add_teacher(teacher_data)
                            teachers_added += 1
                    except Exception as e:
                        errors.append(f"Строка {i}: {str(e)}")
            
            # Обновляем список преподавателей
            self._load_teachers()
            
            # Выводим результат
            message = f"Импорт завершен.\nДобавлено: {teachers_added}\nОбновлено: {teachers_updated}"
            if errors:
                message += f"\n\nОшибки ({len(errors)}):\n" + "\n".join(errors[:10])
                if len(errors) > 10:
                    message += f"\n...и еще {len(errors) - 10} ошибок"
                messagebox.showwarning("Результат импорта", message)
            else:
                messagebox.showinfo("Результат импорта", message)
        except Exception as e:
            logger.error(f"Ошибка при импорте преподавателей: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось импортировать данные: {str(e)}")
    
    def _export_teachers_to_csv(self):
        """Экспорт преподавателей в CSV-файл"""
        file_path = filedialog.asksaveasfilename(
            title="Сохранить CSV-файл",
            defaultextension=".csv",
            filetypes=[("CSV files", "*.csv"), ("All files", "*.*")]
        )
        
        if not file_path:
            return
        
        try:
            import csv
            
            # Получаем всех преподавателей
            teachers = self.app.get_all_teachers()
            
            with open(file_path, 'w', encoding='utf-8', newline='') as csvfile:
                fieldnames = ['id', 'name', 'position', 'academic_degree', 'qualification_category', 
                             'hourly_rate', 'experience_years', 'hire_date', 'birth_date', 
                             'is_young_specialist', 'is_union_member']
                writer = csv.DictWriter(csvfile, fieldnames=fieldnames)
                
                writer.writeheader()
                for teacher in teachers:
                    # Форматируем даты
                    if teacher.get('hire_date'):
                        if isinstance(teacher['hire_date'], str):
                            teacher['hire_date'] = teacher['hire_date']
                        else:
                            teacher['hire_date'] = teacher['hire_date'].strftime("%d.%m.%Y")
                    
                    if teacher.get('birth_date'):
                        if isinstance(teacher['birth_date'], str):
                            teacher['birth_date'] = teacher['birth_date']
                        else:
                            teacher['birth_date'] = teacher['birth_date'].strftime("%d.%m.%Y")
                    
                    writer.writerow(teacher)
            
            messagebox.showinfo("Экспорт", f"Экспортировано преподавателей: {len(teachers)}")
        except Exception as e:
            logger.error(f"Ошибка при экспорте преподавателей: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные: {str(e)}")
    
    # --- Методы для расчета зарплаты ---
    
    def _on_salary_teacher_select(self, event):
        """Обработчик выбора преподавателя на вкладке 'Расчет зарплаты'"""
        selected_name = self.salary_teacher_var.get()
        if not selected_name:
            return
        
        try:
            teachers = self._get_teachers_list()
            teacher_id = None
            for t_id, name in teachers:
                if name == selected_name:
                    teacher_id = t_id
                    break
            
            if teacher_id:
                self.current_teacher_id = teacher_id
                self.update_status(f"Выбран преподаватель для расчета: {selected_name}")
            else:
                self.current_teacher_id = None
                self.update_status("Преподаватель не найден")
        except Exception as e:
            logger.error(f"Ошибка при выборе преподавателя: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось выбрать преподавателя: {str(e)}")
    
    def _display_salary_calculation(self, calculation):
        """
        Отображение результатов расчета зарплаты
        
        :param calculation: словарь с результатами расчета
        """
        try:
            # Получаем имя преподавателя
            teacher_name = calculation.get('teacher_name', '')
            if not teacher_name:
                teacher_id = calculation.get('teacher_id')
                if teacher_id:
                    teacher = self.app.get_teacher_by_id(teacher_id)
                    if teacher:
                        teacher_name = teacher.get('name', f'ID: {teacher_id}')
                    else:
                        teacher_name = f'ID: {teacher_id}'
            
            # Основная информация
            self.result_teacher_var.set(teacher_name)
            
            # Форматируем дату
            calculation_date = calculation.get('calculation_date')
            if isinstance(calculation_date, datetime.date):
                self.result_date_var.set(calculation_date.strftime('%d.%m.%Y'))
            else:
                self.result_date_var.set(str(calculation_date))
            
            # Детали расчета
            self.result_hours_var.set(f"{calculation.get('hours_worked', 0)} ч")
            self.result_rate_var.set(f"{calculation.get('hourly_rate', 0):.2f} руб/ч")
            self.result_base_salary_var.set(f"{calculation.get('base_salary', 0):.2f} руб")
            self.result_position_bonus_var.set(f"{calculation.get('position_bonus', 0):.2f} руб")
            self.result_degree_bonus_var.set(f"{calculation.get('degree_bonus', 0):.2f} руб")
            self.result_experience_bonus_var.set(f"{calculation.get('experience_bonus', 0):.2f} руб")
            self.result_category_bonus_var.set(f"{calculation.get('category_bonus', 0):.2f} руб")
            self.result_young_specialist_bonus_var.set(f"{calculation.get('young_specialist_bonus', 0):.2f} руб")
            self.result_sick_leave_pay_var.set(f"{calculation.get('sick_leave_pay', 0):.2f} руб")
            self.result_bonus_var.set(f"{calculation.get('bonus', 0):.2f} руб")
            
            # Итоги
            self.result_gross_salary_var.set(f"{calculation.get('gross_salary', 0):.2f} руб")
            self.result_tax_rate_var.set(f"{calculation.get('tax_rate', 0)}%")
            self.result_tax_amount_var.set(f"{calculation.get('tax_amount', 0):.2f} руб")
            self.result_union_contribution_var.set(f"{calculation.get('union_contribution', 0):.2f} руб")
            self.result_net_salary_var.set(f"{calculation.get('net_salary', 0):.2f} руб")
        except Exception as e:
            logger.error(f"Ошибка при отображении результатов расчета: {str(e)}")
            self.update_status(f"Ошибка: {str(e)}")

    def _calculate_salary(self):
        """Расчет зарплаты на основе введенных данных"""
        try:
            # Получаем выбранного преподавателя
            selected_name = self.salary_teacher_var.get()
            if not selected_name:
                messagebox.showinfo("Информация", "Выберите преподавателя")
                return
            
            # Получаем ID преподавателя
            teacher_id = self._get_teacher_id_by_name(selected_name)
            if teacher_id is None:
                messagebox.showinfo("Информация", "Преподаватель не найден")
                return
            
            # Получаем данные для расчета
            try:
                hours_worked = float(self.hours_worked_var.get())
                sick_leave_hours = float(self.sick_leave_hours_var.get())
                absence_hours = float(self.absence_hours_var.get())
                bonus = float(self.bonus_var.get())
                tax_rate = float(self.tax_rate_var.get())
                
                # Проверка корректности введенных данных
                if hours_worked < 0 or sick_leave_hours < 0 or absence_hours < 0 or bonus < 0 or tax_rate < 0:
                    messagebox.showwarning("Предупреждение", "Все значения должны быть положительными числами")
                    return
            except ValueError:
                messagebox.showwarning("Предупреждение", "Введите корректные числовые значения")
                return
            
            # Получаем дату расчета
            try:
                calculation_date = datetime.datetime.strptime(self.calc_date_var.get(), "%d.%m.%Y").date()
            except ValueError:
                messagebox.showwarning("Предупреждение", "Неверный формат даты. Используйте ДД.ММ.ГГГГ")
                return
            
            # Формируем данные для расчета
            calc_data = {
                'hours_worked': hours_worked,
                'sick_leave_hours': sick_leave_hours,
                'absence_hours': absence_hours,
                'bonus': bonus,
                'tax_rate': tax_rate,
                'calculation_date': calculation_date
            }
            
            # Выполняем расчет
            result = self.app.calculate_salary(teacher_id, calc_data)
            
            # Сохраняем результат для возможного последующего сохранения
            self.current_calculation = result
            
            # Отображаем результаты
            self._display_salary_calculation(result)
            
            self.update_status("Расчет зарплаты выполнен успешно")
        except Exception as e:
            logger.error(f"Ошибка при расчете зарплаты: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при расчете зарплаты: {str(e)}")
    
    def _save_salary_calculation(self):
        """Сохранить расчет зарплаты в базу данных"""
        if not hasattr(self, 'current_calculation') or not self.current_calculation:
            messagebox.showwarning("Предупреждение", "Нет данных для сохранения. Сначала выполните расчет.")
            return
        
        try:
            # Сохраняем расчет в базу
            calculation_id = self.app.save_salary_calculation(self.current_calculation)
            
            messagebox.showinfo("Успех", f"Расчет заработной платы успешно сохранен (ID: {calculation_id})")
            self.update_status(f"Расчет сохранен (ID: {calculation_id})")
        except Exception as e:
            logger.error(f"Ошибка при сохранении расчета: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось сохранить расчет: {str(e)}")
    
    def _clear_salary_calculation(self):
        """Очистить форму расчета зарплаты"""
        # Очищаем поля ввода
        self.hours_worked_var.set("")
        self.sick_leave_hours_var.set("0")
        self.absence_hours_var.set("0")
        self.bonus_var.set("0")
        self.tax_rate_var.set("13")
        
        # Очищаем результаты
        self.result_teacher_var.set("")
        self.result_date_var.set("")
        self.result_hours_var.set("")
        self.result_rate_var.set("")
        self.result_base_salary_var.set("")
        self.result_position_bonus_var.set("")
        self.result_degree_bonus_var.set("")
        self.result_experience_bonus_var.set("")
        self.result_category_bonus_var.set("")
        self.result_young_specialist_bonus_var.set("")
        self.result_sick_leave_pay_var.set("")
        self.result_bonus_var.set("")
        self.result_gross_salary_var.set("")
        self.result_tax_rate_var.set("")
        self.result_tax_amount_var.set("")
        self.result_union_contribution_var.set("")
        self.result_net_salary_var.set("")
        
        # Сбрасываем сохраненный расчет
        if hasattr(self, 'current_calculation'):
            self.current_calculation = None
        
        self.update_status("Форма расчета очищена")
    
    def _calculate_salary_for_selected(self):
        """Рассчитать зарплату для выбранного преподавателя на вкладке преподавателей"""
        if not self.current_teacher_id:
            messagebox.showinfo("Информация", "Выберите преподавателя")
            return
        
        # Переключаемся на вкладку расчета зарплаты
        self.notebook.select(1)
        
        # Получаем данные преподавателя
        teacher = self.app.get_teacher_by_id(self.current_teacher_id)
        if not teacher:
            return
        
        # Выбираем преподавателя в комбобоксе
        teachers = self._get_teachers_list()
        for _, name in teachers:
            if name == teacher['name']:
                self.salary_teacher_var.set(name)
                break
        
        # Можно автоматически заполнить некоторые поля по умолчанию
        # Например, текущий месяц и год для периода
        today = datetime.date.today()
        self.calc_date_var.set(today.strftime("%d.%m.%Y"))
        
        # Устанавливаем фокус на поле с количеством отработанных часов
        if hasattr(self, 'hours_worked_var'):
            self.hours_worked_var.set("")  # Очищаем поле
    
    def _export_salary_data(self, teacher_id, start_date, end_date, include_chart=False, output_format='pdf'):
        """
        Экспортирует данные о зарплате преподавателя за указанный период
        
        :param teacher_id: ID преподавателя
        :param start_date: начальная дата периода (строка в формате дд.мм.гггг или объект datetime)
        :param end_date: конечная дата периода (строка в формате дд.мм.гггг или объект datetime)
        :param include_chart: включать ли график в отчет
        :param output_format: формат выходного файла ('pdf', 'csv')
        :return: путь к созданному файлу, если экспорт успешен, иначе None
        """
        try:
            logger.info(f"Экспорт данных о зарплате для преподавателя ID={teacher_id} за период {start_date} - {end_date}")
            
            # Получаем данные преподавателя
            teacher_info = self.app.get_teacher_by_id(teacher_id)
            if not teacher_info:
                raise ValueError(f"Преподаватель с ID={teacher_id} не найден")
            
            # Получаем данные о зарплате за период
            salary_data = self.app.get_salary_data_for_period(teacher_id, start_date, end_date)
            
            if not salary_data:
                messagebox.showinfo("Информация", "Нет данных о зарплате за указанный период")
                return None
            
            # Предлагаем пользователю выбрать место для сохранения
            if output_format.lower() == 'pdf':
                file_types = [("PDF файлы", "*.pdf"), ("Все файлы", "*.*")]
                default_ext = ".pdf"
            elif output_format.lower() == 'csv':
                file_types = [("CSV файлы", "*.csv"), ("Все файлы", "*.*")]
                default_ext = ".csv"
            else:
                file_types = [("Все файлы", "*.*")]
                default_ext = ""
                
            output_file = filedialog.asksaveasfilename(
                defaultextension=default_ext,
                filetypes=file_types,
                title="Сохранить отчет как...",
                initialfile=f"salary_report_{teacher_info['name']}_{start_date}_{end_date}{default_ext}"
            )
            
            if not output_file:  # Пользователь отменил сохранение
                return None
            
            # Экспорт в зависимости от формата
            if output_format.lower() == 'pdf':
                # Создаем PDF отчет
                if not self._create_pdf_report(
                    salary_data=salary_data,
                    teacher_info=teacher_info,
                    period_start=start_date,
                    period_end=end_date,
                    output_file=output_file,
                    include_chart=include_chart
                ):
                    return None
            elif output_format.lower() == 'csv':
                # Экспорт в CSV (можно добавить отдельный метод для этого)
                with open(output_file, 'w', newline='', encoding='utf-8') as csvfile:
                    import csv
                    writer = csv.writer(csvfile)
                    # Заголовки
                    writer.writerow(['Дата', 'Отработано часов', 'Ставка', 'Бонус', 'Налог', 'Итого'])
                    # Данные
                    for entry in salary_data:
                        writer.writerow([
                            entry.get('calculation_date', ''),
                            entry.get('hours_worked', 0),
                            entry.get('hourly_rate', 0),
                            entry.get('bonus_amount', 0),
                            entry.get('tax_amount', 0),
                            entry.get('net_salary', 0)
                        ])
            else:
                messagebox.showerror("Ошибка", f"Неподдерживаемый формат файла: {output_format}")
                return None
                
            messagebox.showinfo("Успех", f"Отчет успешно сохранен в файл:\n{output_file}")
            return output_file
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных о зарплате: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные о зарплате: {str(e)}")
            return None


    def _clean_numeric_value(self, value_str, suffix=None):
        """
        Очищает строковое значение от суффикса и возвращает числовое значение.
        
        :param value_str: Строковое представление числа, возможно с суффиксом
        :param suffix: Суффикс для удаления (если None, удаляются все нецифровые символы кроме точки)
        :return: Числовое значение (float)
        """
        if not value_str:
            return 0.0
        
        # Приводим к строке для безопасности
        value_str = str(value_str)
        
        try:
            if suffix:
                # Удаляем конкретный суффикс
                clean_str = value_str.replace(suffix, '').strip()
            else:
                # Оставляем только цифры, точку и знак минуса
                import re
                clean_str = re.sub(r'[^\d.-]', '', value_str)
                
            return float(clean_str)
        except (ValueError, TypeError) as e:
            logger.warning(f"Не удалось преобразовать '{value_str}' в число: {e}")
            return 0.0

    def _export_salary_to_pdf(self):
        """Экспорт расчета зарплаты в PDF"""
        try:
            if not hasattr(self, 'current_calculation') or not self.current_calculation:
                messagebox.showerror("Ошибка", "Нет данных для экспорта. Сначала выполните расчет.")
                return
            
            # Запрашиваем путь для сохранения файла
            file_path = filedialog.asksaveasfilename(
                defaultextension=".pdf",
                filetypes=[("PDF файлы", "*.pdf")],
                title="Сохранить отчет как"
            )
            
            if not file_path:
                return  # Пользователь отменил сохранение
            
            # Проверяем наличие необходимых библиотек
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib.units import mm
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
            except ImportError:
                messagebox.showerror(
                    "Отсутствуют зависимости", 
                    "Для экспорта в PDF необходимы дополнительные библиотеки.\n"
                    "Установите их командой: pip install reportlab"
                )
                return
            
            # Регистрируем шрифты для поддержки кириллицы
            # Используем один из стандартных шрифтов с кириллицей
            use_default_fonts = False
            try:
                # Попытка зарегистрировать шрифт Arial или DejaVuSans
                fonts_to_try = [
                    ('Arial', 'arial.ttf'),
                    ('DejaVuSans', 'DejaVuSans.ttf'),
                    ('Times New Roman', 'times.ttf'),
                    ('Verdana', 'verdana.ttf')
                ]
                
                for font_name, font_file in fonts_to_try:
                    try:
                        pdfmetrics.registerFont(TTFont(font_name, font_file))
                        logger.info(f"Зарегистрирован шрифт: {font_name}")
                        break
                    except Exception as e:
                        logger.warning(f"Не удалось зарегистрировать шрифт {font_name}: {str(e)}")
                else:  # Если не удалось зарегистрировать ни один шрифт
                    use_default_fonts = True
                    logger.warning("Не удалось зарегистрировать ни один из шрифтов с поддержкой кириллицы")
            except Exception as e:
                use_default_fonts = True
                logger.warning(f"Ошибка при регистрации шрифтов: {str(e)}")
            
            # Начинаем создавать PDF
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # Создаем стили для кириллических текстов
            if not use_default_fonts:
                # Получаем имя зарегистрированного шрифта
                registered_fonts = pdfmetrics.getRegisteredFontNames()
                cyrillic_font = 'Helvetica'  # Шрифт по умолчанию, если ни один не зарегистрирован
                
                for font in ['Arial', 'DejaVuSans', 'Times New Roman', 'Verdana']:
                    if font in registered_fonts:
                        cyrillic_font = font
                        break
                
                # Создаем стиль для кириллицы
                cyrillic_style = ParagraphStyle(
                    name='Cyrillic',
                    fontName=cyrillic_font,
                    fontSize=10,
                    leading=12,
                    encoding='utf-8'
                )
                
                heading_style = ParagraphStyle(
                    name='CyrillicHeading',
                    fontName=cyrillic_font,
                    fontSize=14,
                    leading=16,
                    alignment=1,  # По центру
                    spaceAfter=10,
                    spaceBefore=10,
                    encoding='utf-8',
                    bold=True
                )
            else:
                # Используем стандартные стили, но с кодировкой UTF-8
                cyrillic_style = styles['Normal']
                cyrillic_style.encoding = 'utf-8'
                
                heading_style = styles['Heading2']
                heading_style.encoding = 'utf-8'
            
            # Формируем содержимое документа
            content = []
            
            # Заголовок отчета (используем кириллический стиль)
            title = f"Расчет заработной платы для {self.current_calculation.get('teacher_name', '')}"
            content.append(Paragraph(title, heading_style))
            content.append(Spacer(1, 10 * mm))
            
            # Дата расчета (используем кириллический стиль)
            calc_date = self.current_calculation.get('calculation_date', datetime.datetime.now()).strftime("%d.%m.%Y")
            content.append(Paragraph(f"Дата расчета: {calc_date}", cyrillic_style))
            content.append(Spacer(1, 5 * mm))
            
            # Безопасно получаем значения из словаря с проверкой на существование ключей
            def safe_get(dictionary, key, default=0.0):
                """Безопасно получает значение из словаря с проверкой на существование ключа"""
                return dictionary.get(key, default)
            
            # Подготавливаем данные для таблицы
            # Все строковые значения должны быть в кодировке UTF-8
            table_data = [
                ["Параметр", "Значение"],
                ["Отработано часов", f"{safe_get(self.current_calculation, 'hours_worked')}"],
                ["Часовая ставка", f"{safe_get(self.current_calculation, 'hourly_rate'):.2f} руб/ч"],
                ["Базовая зарплата", f"{safe_get(self.current_calculation, 'base_salary'):.2f} руб"],
                ["Надбавка за должность", f"{safe_get(self.current_calculation, 'position_bonus'):.2f} руб"],
                ["Надбавка за степень", f"{safe_get(self.current_calculation, 'degree_bonus'):.2f} руб"],
                ["Надбавка за стаж", f"{safe_get(self.current_calculation, 'experience_bonus'):.2f} руб"],
                ["Надбавка за категорию", f"{safe_get(self.current_calculation, 'category_bonus'):.2f} руб"],
                ["Надбавка молодому специалисту", f"{safe_get(self.current_calculation, 'young_specialist_bonus'):.2f} руб"],
                ["Оплата больничных", f"{safe_get(self.current_calculation, 'sick_leave_payment'):.2f} руб"],
                ["Бонус", f"{safe_get(self.current_calculation, 'bonus'):.2f} руб"],
                ["Валовая зарплата", f"{safe_get(self.current_calculation, 'gross_salary'):.2f} руб"],
                ["Ставка налога", f"{safe_get(self.current_calculation, 'tax_rate'):.2f}%"],
                ["Сумма налога", f"{safe_get(self.current_calculation, 'tax_amount'):.2f} руб"],
                ["Профсоюзные взносы", f"{safe_get(self.current_calculation, 'union_contribution'):.2f} руб"],
                ["ЧИСТАЯ ЗАРПЛАТА", f"{safe_get(self.current_calculation, 'net_salary'):.2f} руб"]
            ]
            
            # Преобразуем данные таблицы в Paragraph с правильным стилем для кириллицы
            formatted_data = []
            for row in table_data:
                formatted_row = []
                for cell in row:
                    # Создаем Paragraph для каждой ячейки, чтобы применить кириллический шрифт
                    formatted_row.append(Paragraph(str(cell), cyrillic_style))
                formatted_data.append(formatted_row)
            
            # Создаем таблицу с форматированными данными
            table = Table(formatted_data, colWidths=[100 * mm, 70 * mm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (1, 0), 'CENTER'),
                ('FONTNAME', (0, 0), (1, 0), cyrillic_style.fontName),
                ('FONTSIZE', (0, 0), (1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (1, 0), 12),
                ('BACKGROUND', (0, -1), (1, -1), colors.lightgrey),
                ('FONTNAME', (0, -1), (1, -1), cyrillic_style.fontName),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ]))
            
            content.append(table)
            content.append(Spacer(1, 10 * mm))
            
            # Примечания (используем кириллический стиль)
            content.append(Paragraph("Примечания:", cyrillic_style))
            content.append(Paragraph("1. Расчет выполнен в соответствии с действующим законодательством.", cyrillic_style))
            content.append(Paragraph("2. Для получения дополнительной информации обратитесь в бухгалтерию.", cyrillic_style))
            
            # Формируем PDF
            doc.build(content)
            
            # Обновляем статус и уведомляем пользователя
            self.update_status(f"Расчет зарплаты экспортирован в PDF: {file_path}")
            messagebox.showinfo("Экспорт завершен", f"Отчет сохранен в файл:\n{file_path}")
            
            # Открываем файл в ассоциированной программе
            try:
                import os
                import platform
                if platform.system() == 'Windows':
                    os.startfile(file_path)
                elif platform.system() == 'Darwin':  # macOS
                    import subprocess
                    subprocess.call(('open', file_path))
                else:  # Linux
                    import subprocess
                    subprocess.call(('xdg-open', file_path))
            except Exception as e:
                logger.error(f"Не удалось открыть PDF-файл: {str(e)}")
            
            return file_path
        
        except Exception as e:
            logger.error(f"Ошибка при экспорте расчета зарплаты: {str(e)}")
            messagebox.showerror("Ошибка", f"Ошибка при экспорте расчета зарплаты: {str(e)}")
            return None


    
    def _print_salary_result(self):
        """Печать результатов расчета зарплаты"""
        if not hasattr(self, 'current_calculation') or not self.current_calculation:
            messagebox.showwarning("Предупреждение", "Нет данных для печати. Сначала выполните расчет.")
            return
        
        try:
            # В простейшем случае можно сначала экспортировать в PDF, а затем открыть этот файл
            import tempfile
            import os
            import subprocess
            
            # Создаем временный файл
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                temp_path = tmp.name
            
            # Экспортируем в этот файл
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            except ImportError:
                messagebox.showerror(
                    "Ошибка импорта", 
                    "Отсутствует пакет reportlab. Установите его командой:\npip install reportlab"
                )
                return
            
            calc = self.current_calculation
            
            # Создаем PDF документ
            doc = SimpleDocTemplate(temp_path, pagesize=A4)
            
            # Стили и содержимое как в методе экспорта в PDF
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(name='Russian', fontName='Helvetica', fontSize=12, leading=14))
            styles.add(ParagraphStyle(name='RussianBold', fontName='Helvetica-Bold', fontSize=14, leading=16))
            styles.add(ParagraphStyle(name='RussianTitle', fontName='Helvetica-Bold', fontSize=18, leading=22, alignment=1))
            
            content = []
            content.append(Paragraph("РАСЧЕТНЫЙ ЛИСТ", styles['RussianTitle']))
            content.append(Spacer(1, 12))
            content.append(Paragraph(f"Преподаватель: {calc['teacher_name']}", styles['RussianBold']))
            content.append(Paragraph(f"Дата расчета: {calc['calculation_date'].strftime('%d.%m.%Y')}", styles['Russian']))
            content.append(Spacer(1, 12))
            
            # Таблица с данными (аналогично _export_salary_to_pdf)
            data = [
                ["Параметр", "Значение"],
                ["Отработано часов", f"{calc['hours_worked']}"],
                ["Часовая ставка", f"{calc['hourly_rate']:.2f} руб/ч"],
                ["Базовая зарплата", f"{calc['base_salary']:.2f} руб"],
                ["Надбавка за должность", f"{calc['position_bonus']:.2f} руб"],
                ["Надбавка за ученую степень", f"{calc['degree_bonus']:.2f} руб"],
                ["Надбавка за стаж", f"{calc['experience_bonus']:.2f} руб"],
                ["Надбавка за категорию", f"{calc['category_bonus']:.2f} руб"]
            ]
            
            if calc.get('young_specialist_bonus', 0) > 0:
                data.append(["Надбавка молодому специалисту", f"{calc['young_specialist_bonus']:.2f} руб"])
            
            if calc.get('sick_leave_pay', 0) > 0:
                data.append(["Оплата больничных", f"{calc['sick_leave_pay']:.2f} руб"])
            
            if calc.get('bonus', 0) > 0:
                data.append(["Бонус", f"{calc['bonus']:.2f} руб"])
            
            data.append(["", ""])
            data.append(["ВАЛОВАЯ ЗАРПЛАТА", f"{calc['gross_salary']:.2f} руб"])
            data.append(["Подоходный налог", f"{calc.get('tax_amount', 0):.2f} руб"])
            
            if calc.get('union_contribution', 0) > 0:
                data.append(["Профсоюзные взносы", f"{calc['union_contribution']:.2f} руб"])
            
            data.append(["", ""])
            data.append(["ЧИСТАЯ ЗАРПЛАТА К ВЫПЛАТЕ", f"{calc['net_salary']:.2f} руб"])
            
            table = Table(data, colWidths=[300, 150])
            table_style = TableStyle([
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
                ('ALIGN', (1, 1), (1, -1), 'RIGHT'),
                ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
                ('LINEBELOW', (0, 0), (-1, 0), 1, colors.black),
                ('LINEABOVE', (0, -1), (-1, -1), 1, colors.black),
                ('GRID', (0, 0), (-1, -2), 0.5, colors.grey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('PADDING', (0, 0), (-1, -1), 6),
            ])
            table.setStyle(table_style)
            
            content.append(table)
            content.append(Spacer(1, 20))
            
            current_date = datetime.datetime.now().strftime('%d.%m.%Y %H:%M')
            content.append(Paragraph(f"Документ сформирован: {current_date}", styles['Russian']))
            
            # Создаем PDF
            doc.build(content)
            
            # Открываем PDF в ассоциированной программе
            if os.name == 'nt':  # Windows
                os.startfile(temp_path)
            elif os.name == 'posix':  # Linux/Mac
                subprocess.call(('xdg-open', temp_path))
            
            # Временный файл будет удален, когда пользователь закроет программу просмотра PDF
            self.update_status("Документ отправлен на печать")
        except Exception as e:
            logger.error(f"Ошибка при печати: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось выполнить печать: {str(e)}")
    
    # --- Методы для работы с отпусками ---
    
    def _on_vacation_select(self, event):

        try:
            selected_name = self.salary_teacher_var.get()
            if not selected_name:
                return
            
            # Получение ID преподавателя по имени с помощью вспомогательного метода
            teacher_id = self._get_teacher_id_by_name(selected_name)
            
            if teacher_id is not None:
                # Получаем данные преподавателя
                teacher = self.app.get_teacher_by_id(teacher_id)
                if not teacher:
                    return
                    
                # Заполняем информацию о ставке или другие необходимые данные
                hourly_rate = teacher.get('hourly_rate', 0)
                if hourly_rate:
                    self.update_status(f"Выбран преподаватель: {selected_name}, ставка: {hourly_rate} руб/час")
                
                # Здесь можно добавить дополнительные действия при выборе преподавателя
                # Например, заполнить поля с данными для расчета
        except Exception as e:
            logger.error(f"Ошибка при выборе преподавателя для расчета: {str(e)}")
            self.update_status(f"Ошибка: {str(e)}")

        """Обработчик выбора отпуска в таблице"""
        selected_items = self.vacations_tree.selection()
        if not selected_items:
            return
        
        item = selected_items[0]
        vacation_id = self.vacations_tree.item(item, 'values')[0]
        
        try:
            # Сохраняем ID текущего выбранного отпуска
            self.current_vacation_id = int(vacation_id)
            
            # Дополнительные действия при необходимости
            self.update_status(f"Выбран отпуск (ID: {vacation_id})")
        except Exception as e:
            logger.error(f"Ошибка при выборе отпуска: {str(e)}")
    
    def _schedule_vacation(self):
        """Показать диалог планирования отпуска"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель")
            return
        
        # Просто показываем правую панель с формой планирования
        self.update_status("Используйте правую панель для планирования отпуска")
        
        # Подготавливаем форму
        self.vacation_start_date_var.set("")
        self.vacation_end_date_var.set("")
        self.vacation_notes_var.set("")
        self.vacation_type_var.set("основной")
    
    def _schedule_new_vacation(self):
        """Запланировать новый отпуск по данным из формы"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель")
            return
        
        try:
            # Проверка обязательных полей
            if not self.vacation_start_date_var.get().strip():
                messagebox.showerror("Ошибка", "Поле 'Дата начала' обязательно для заполнения")
                return
            
            if not self.vacation_end_date_var.get().strip():
                messagebox.showerror("Ошибка", "Поле 'Дата окончания' обязательно для заполнения")
                return
            
            # Преобразование дат
            start_date = datetime.datetime.strptime(self.vacation_start_date_var.get(), "%d.%m.%Y").date()
            end_date = datetime.datetime.strptime(self.vacation_end_date_var.get(), "%d.%m.%Y").date()
            
            # Проверка корректности дат
            if start_date > end_date:
                messagebox.showerror("Ошибка", "Дата начала не может быть позже даты окончания")
                return
            
            # Планируем отпуск
            vacation_id = self.app.schedule_vacation(
                self.current_teacher_id,
                start_date,
                end_date,
                self.vacation_type_var.get(),
                self.vacation_notes_var.get()
            )
            
            # Обновляем список отпусков
            self._load_vacations()
            
            # Обновляем доступные дни отпуска
            available_days = self.app.get_teacher_remaining_vacation_days(self.current_teacher_id)
            self.available_days_var.set(str(available_days))
            
            # Очищаем форму
            self.vacation_start_date_var.set("")
            self.vacation_end_date_var.set("")
            self.vacation_notes_var.set("")
            
            messagebox.showinfo("Успех", f"Отпуск успешно запланирован (ID: {vacation_id})")
        except Exception as e:
            logger.error(f"Ошибка при планировании отпуска: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось запланировать отпуск: {str(e)}")
    
    def _cancel_vacation(self):
        """Отменить выбранный отпуск"""
        if not hasattr(self, 'current_vacation_id') or not self.current_vacation_id:
            messagebox.showwarning("Предупреждение", "Не выбран отпуск для отмены")
            return
        
        try:
            # Запрашиваем подтверждение
            confirm = messagebox.askyesno(
                "Подтверждение", 
                "Вы действительно хотите отменить выбранный отпуск?"
            )
            
            if not confirm:
                return
            
            # Отменяем отпуск
            success = self.app.cancel_vacation(self.current_vacation_id)
            
            if success:
                # Обновляем список отпусков
                self._load_vacations()
                
                # Обновляем доступные дни отпуска
                available_days = self.app.get_teacher_remaining_vacation_days(self.current_teacher_id)
                self.available_days_var.set(str(available_days))
                
                messagebox.showinfo("Успех", "Отпуск успешно отменен")
            else:
                messagebox.showwarning("Предупреждение", "Не удалось отменить отпуск. Возможно, он уже имеет статус, не позволяющий отмену.")
        except Exception as e:
            logger.error(f"Ошибка при отмене отпуска: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось отменить отпуск: {str(e)}")
    
    def _mark_vacation_as_used(self):
        """Отметить выбранный отпуск как использованный"""
        if not hasattr(self, 'current_vacation_id') or not self.current_vacation_id:
            messagebox.showwarning("Предупреждение", "Не выбран отпуск")
            return
        
        try:
            # Запрашиваем подтверждение
            confirm = messagebox.askyesno(
                "Подтверждение", 
                "Вы действительно хотите отметить выбранный отпуск как использованный?"
            )
            
            if not confirm:
                return
            
            # Отмечаем отпуск как использованный
            success = self.app.vacation_processor.mark_vacation_as_used(self.current_vacation_id)
            
            if success:
                # Обновляем список отпусков
                self._load_vacations()
                
                messagebox.showinfo("Успех", "Отпуск отмечен как использованный")
            else:
                messagebox.showwarning("Предупреждение", "Не удалось изменить статус отпуска. Возможно, он уже имеет другой статус.")
        except Exception as e:
            logger.error(f"Ошибка при изменении статуса отпуска: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось изменить статус отпуска: {str(e)}")
    
    def _calculate_vacation_payment(self, vacation_id=None):
        """Расчет выплаты за отпуск"""
        # Переключаемся на вкладку отпусков
        self.notebook.select(2)
        
        # Адаптер для вызова расчета выплат за отпуск
        if hasattr(self, 'vacation_tab_manager') and hasattr(self.vacation_tab_manager, '_calculate_vacation_payment'):
            # Если ID отпуска не передан, то пусть класс VacationTab сам выберет отпуск
            if vacation_id is None:
                if hasattr(self.vacation_tab_manager, '_get_selected_vacation_id'):
                    vacation_id = self.vacation_tab_manager._get_selected_vacation_id()
            
            # Если ID отпуска определен, вызываем расчет
            if vacation_id:
                self.vacation_tab_manager._calculate_vacation_payment(vacation_id)
            else:
                messagebox.showwarning("Предупреждение", "Сначала выберите отпуск для расчета")
        else:
            messagebox.showwarning("Предупреждение", "Функциональность расчета отпускных недоступна")
    
    def _calculate_optimal_vacation(self):
        """Рассчитать оптимальное распределение отпуска"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель")
            return
        
        try:
            # Получаем год из поля ввода
            try:
                year = int(self.optimal_year_var.get())
            except ValueError:
                messagebox.showerror("Ошибка", "Некорректно указан год")
                return
            
            # Получаем предложения по отпуску
            suggestions = self.app.suggest_optimal_vacation_distribution(self.current_teacher_id, year)
            
            # Очищаем текстовое поле
            self.optimal_suggestions_text.delete(1.0, tk.END)
            
            if not suggestions:
                self.optimal_suggestions_text.insert(tk.END, "Не удалось сгенерировать предложения по отпуску.")
                return
            
            # Выводим предложения
            for i, suggestion in enumerate(suggestions, start=1):
                self.optimal_suggestions_text.insert(tk.END, f"{i}. {suggestion['description']}\n\n")
            
            self.update_status(f"Сгенерировано {len(suggestions)} предложений по отпуску")
        except Exception as e:
            logger.error(f"Ошибка при расчете оптимального распределения отпуска: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось рассчитать оптимальное распределение отпуска: {str(e)}")
    
    def _apply_selected_suggestion(self):
        """Применить выбранное предложение по отпуску"""
        if not self.current_teacher_id:
            messagebox.showwarning("Предупреждение", "Не выбран преподаватель")
            return
        
        try:
            # Получаем выбранное предложение
            selected_text = self.optimal_suggestions_text.get(tk.SEL_FIRST, tk.SEL_LAST)
            if not selected_text:
                messagebox.showwarning("Предупреждение", "Не выбрано предложение по отпуску")
                return
            
            # Извлекаем даты из предложения
            import re
            match = re.search(r'с (\d{2}\.\d{2}\.\d{4}) по (\d{2}\.\d{2}\.\d{4})', selected_text)
            if not match:
                messagebox.showerror("Ошибка", "Не удалось извлечь даты из выбранного предложения")
                return
            
            start_date_str, end_date_str = match.groups()
            
            # Устанавливаем даты в форму планирования
            self.vacation_start_date_var.set(start_date_str)
            self.vacation_end_date_var.set(end_date_str)
            
            messagebox.showinfo("Информация", "Даты установлены в форму планирования отпуска. Нажмите 'Запланировать отпуск' для подтверждения.")
        except Exception as e:
            logger.error(f"Ошибка при применении предложения по отпуску: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось применить предложение: {str(e)}")
    
    def _schedule_vacation_for_selected(self):
        """Открывает вкладку отпусков для выбранного преподавателя"""
        if self.current_teacher_id:
            # Переключаемся на вкладку отпусков
            self.notebook.select(2)  # Индекс вкладки отпусков (считая от 0)
            
            # Используем функциональность класса VacationTab для выбора преподавателя
            if hasattr(self, 'vacation_tab_manager'):
                teacher = self.app.get_teacher_by_id(self.current_teacher_id)
                if teacher:
                    # Находим имя преподавателя в списке
                    teacher_name = teacher.get('name', '')
                    if teacher_name and hasattr(self.vacation_tab_manager, 'teacher_var'):
                        self.vacation_tab_manager.teacher_var.set(teacher_name)
                        # Вызываем обработчик выбора преподавателя
                        if hasattr(self.vacation_tab_manager, '_on_teacher_selected'):
                            self.vacation_tab_manager._on_teacher_selected(None)
        else:
            messagebox.showwarning("Предупреждение", "Сначала выберите преподавателя")

    def _generate_vacation_report(self):
        """Генерирует отчет по отпускам на основе выбранных параметров"""
        # Здесь будет код для генерации отчета по отпускам,
        # можно перенести существующий функционал из vacation_tab_manager._export_vacation_report()
        messagebox.showinfo("Информация", "Функция генерации отчета по отпускам находится в разработке")


    def _get_teacher_id_by_name(self, teacher_name):
        """
        Получает ID преподавателя по его имени
        
        :param teacher_name: имя преподавателя
        :return: ID преподавателя или None, если преподаватель не найден
        """
        try:
            teachers = self._get_teachers_list()
            for teacher_id, name in teachers:
                if name == teacher_name:
                    return teacher_id
            return None
        except Exception as e:
            logger.error(f"Ошибка при получении ID преподавателя: {str(e)}")
            return None
    
    def _prepare_salary_data_for_export(self, teacher, salary_data):
        """
        Подготовка данных о зарплате для экспорта
        
        :param teacher: Словарь с информацией о преподавателе
        :param salary_data: Необработанные данные о расчетах зарплаты
        :return: Отформатированные данные, готовые для экспорта
        """
        # Форматируем данные для экспорта
        formatted_data = {
            'teacher': teacher,
            'calculations': salary_data,
            'summary': {
                'total_gross': sum(calc['gross_salary'] for calc in salary_data),
                'total_net': sum(calc['net_salary'] for calc in salary_data),
                'total_hours': sum(calc['hours_worked'] for calc in salary_data),
                'period_start': min(calc['calculation_date'] for calc in salary_data),
                'period_end': max(calc['calculation_date'] for calc in salary_data)
            }
        }
        
        return formatted_data


    def _generate_salary_report(self):
        """Генерация отчета по зарплате выбранного преподавателя за период"""
        try:
            # Проверка выбора преподавателя
            if not hasattr(self, 'report_teacher_var') or not self.report_teacher_var.get():
                messagebox.showwarning("Предупреждение", "Выберите преподавателя для отчета")
                return
                
            teacher_name = self.report_teacher_var.get()
            if not teacher_name or teacher_name not in self.teachers_mapping:
                messagebox.showwarning("Предупреждение", "Выберите корректного преподавателя")
                return
                
            teacher_id = self.teachers_mapping[teacher_name]
            
            # Получаем период отчета
            try:
                start_date = self.start_date_var.get()
                end_date = self.end_date_var.get()
                
                # Проверяем формат дат
                datetime.datetime.strptime(start_date, "%d.%m.%Y")
                datetime.datetime.strptime(end_date, "%d.%m.%Y")
            except ValueError:
                messagebox.showwarning("Предупреждение", "Введите корректные даты в формате ДД.ММ.ГГГГ")
                return
                
            # Получаем выбранные опции
            include_chart = self.report_include_chart_var.get() if hasattr(self, 'report_include_chart_var') else False
            
            # Вызываем экспорт данных
            output_file = self._export_salary_data(
                teacher_id=teacher_id,
                start_date=start_date,
                end_date=end_date,
                include_chart=include_chart
            )
            
            self.update_status(f"Отчет по зарплате для {teacher_name} сгенерирован")
            
        except Exception as e:
            logger.error(f"Ошибка при генерации отчета по зарплате: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось сгенерировать отчет: {str(e)}")


    def _export_salary_summary(self, start_date=None, end_date=None, output_format='pdf', title='Сводный отчет по зарплате', include_details=True, include_chart=True):
        """
        Экспортирует сводные данные о зарплате всех преподавателей за указанный период в файл.
        
        :param start_date: Начальная дата периода (datetime.date)
        :param end_date: Конечная дата периода (datetime.date)
        :param output_format: Формат выходного файла ('pdf', 'excel', 'csv')
        :param title: Заголовок отчета
        :param include_details: Включать ли детали расчета
        :param include_chart: Включать ли графики
        :return: Путь к созданному файлу или None в случае ошибки
        """
        try:
            # Проверяем даты
            if start_date is None:
                start_date = date.today().replace(day=1)  # Начало текущего месяца
                logger.info(f"Не указана начальная дата, используем начало текущего месяца: {start_date}")
                
            if end_date is None:
                end_date = date.today()  # Текущая дата
                logger.info(f"Не указана конечная дата, используем текущую дату: {end_date}")
                
            if isinstance(start_date, str):
                try:
                    start_date = datetime.datetime.strptime(start_date, "%d.%m.%Y").date()
                except ValueError:
                    logger.error(f"Неверный формат начальной даты: {start_date}")
                    raise ValueError(f"Неверный формат начальной даты: {start_date}. Используйте ДД.ММ.ГГГГ")
                    
            if isinstance(end_date, str):
                try:
                    end_date = datetime.datetime.strptime(end_date, "%d.%m.%Y").date()
                except ValueError:
                    logger.error(f"Неверный формат конечной даты: {end_date}")
                    raise ValueError(f"Неверный формат конечной даты: {end_date}. Используйте ДД.ММ.ГГГГ")
            
            if start_date > end_date:
                error_msg = "Начальная дата не может быть позже конечной даты"
                logger.error(error_msg)
                raise ValueError(error_msg)
            
            logger.info(f"Экспорт сводных данных о зарплате за период {start_date} - {end_date}")
            
            # Получаем список всех преподавателей
            teachers = self.app.get_all_teachers()
            if not teachers:
                error_msg = "Не найдено ни одного преподавателя"
                logger.error(error_msg)
                raise ValueError(error_msg)
            
            # Получаем данные о зарплате для всех преподавателей
            all_salary_data = []
            for teacher in teachers:
                teacher_salary_data = self.app.get_salary_data_for_period(
                    teacher_id=teacher['id'],
                    start_date=start_date,
                    end_date=end_date
                )
                
                # Добавляем информацию о преподавателе к каждой записи
                for data in teacher_salary_data:
                    data['teacher'] = teacher
                
                all_salary_data.extend(teacher_salary_data)
            
            if not all_salary_data:
                logger.warning(f"Нет данных о зарплате за период {start_date} - {end_date}")
                messagebox.showinfo("Информация", "Нет данных о зарплате за указанный период")
                return None
            
            # Формируем имя файла
            date_str = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_dir = os.path.join(os.getcwd(), "reports")
            os.makedirs(output_dir, exist_ok=True)
            
            period_str = f"{start_date.strftime('%Y%m%d')}_to_{end_date.strftime('%Y%m%d')}"
            filename_base = f"salary_summary_report_{period_str}_{date_str}"
            
            # Создаем отчет в зависимости от формата
            if output_format.lower() == 'pdf':
                output_file = os.path.join(output_dir, f"{filename_base}.pdf")
                if not self._create_summary_pdf_report(
                    output_file=output_file,
                    salary_data=all_salary_data,
                    teachers=teachers,
                    start_date=start_date,
                    end_date=end_date,
                    title=title,
                    include_details=include_details,
                    include_chart=include_chart
                ):
                    return None
            elif output_format.lower() == 'excel':
                output_file = os.path.join(output_dir, f"{filename_base}.xlsx")
                if not self._create_summary_excel_report(
                    output_file=output_file,
                    salary_data=all_salary_data,
                    teachers=teachers,
                    start_date=start_date,
                    end_date=end_date,
                    title=title,
                    include_details=include_details,
                    include_chart=include_chart
                ):
                    return None
            elif output_format.lower() == 'csv':
                output_file = os.path.join(output_dir, f"{filename_base}.csv")
                if not self._create_summary_csv_report(
                    output_file=output_file,
                    salary_data=all_salary_data,
                    teachers=teachers,
                    start_date=start_date,
                    end_date=end_date
                ):
                    return None
            else:
                logger.error(f"Неподдерживаемый формат отчета: {output_format}")
                messagebox.showerror("Ошибка", f"Неподдерживаемый формат отчета: {output_format}")
                return None
            
            logger.info(f"Сводный отчет успешно создан: {output_file}")
            return output_file
        except ValueError as e:
            # Для ошибок валидации показываем сообщение пользователю
            logger.error(f"Ошибка при экспорте сводных данных о зарплате: {str(e)}")
            messagebox.showerror("Ошибка", str(e))
            return None
        except Exception as e:
            # Для других ошибок логируем подробности
            logger.error(f"Ошибка при экспорте сводных данных о зарплате: {str(e)}", exc_info=True)
            messagebox.showerror("Ошибка", f"Ошибка при экспорте сводных данных о зарплате: {str(e)}")
            return None


    def _export_vacation_data(self, data, format_type, title, include_chart=True, is_summary=False):
        """
        Экспорт данных об отпусках в различные форматы
        
        :param data: Данные об отпусках
        :param format_type: Формат экспорта (PDF, Excel, Word)
        :param title: Заголовок отчета
        :param include_chart: Включать ли график отпусков
        :param is_summary: Является ли отчет сводным по всем преподавателям
        :return: Путь к созданному файлу
        """
        try:
            # Определяем путь для сохранения отчета
            reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            file_path = None
            
            # Экспорт в зависимости от формата
            if format_type == "PDF":
                file_path = os.path.join(reports_dir, f"vacation_report_{timestamp}.pdf")
                self._export_vacation_to_pdf(data, file_path, title, include_chart, is_summary)
                
            elif format_type == "Excel":
                file_path = os.path.join(reports_dir, f"vacation_report_{timestamp}.xlsx")
                self._export_vacation_to_excel(data, file_path, title, include_chart, is_summary)
                
            elif format_type == "Word":
                file_path = os.path.join(reports_dir, f"vacation_report_{timestamp}.docx")
                self._export_vacation_to_word(data, file_path, title, include_chart, is_summary)
            
            logger.info(f"Отчет об отпусках успешно экспортирован в {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных об отпусках: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные: {str(e)}")
            return None

    def _export_vacation_calendar(self, data, format_type, title, include_chart=True):
        """
        Экспорт календарного графика отпусков
        
        :param data: Данные для календарного графика
        :param format_type: Формат экспорта (PDF, Excel, Word)
        :param title: Заголовок отчета
        :param include_chart: Включать ли визуальный график
        :return: Путь к созданному файлу
        """
        try:
            # Определяем путь для сохранения отчета
            reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
            if not os.path.exists(reports_dir):
                os.makedirs(reports_dir)
            
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            file_path = None
            
            # Экспорт в зависимости от формата
            if format_type == "PDF":
                file_path = os.path.join(reports_dir, f"vacation_calendar_{timestamp}.pdf")
                self._export_calendar_to_pdf(data, file_path, title, include_chart)
                
            elif format_type == "Excel":
                file_path = os.path.join(reports_dir, f"vacation_calendar_{timestamp}.xlsx")
                self._export_calendar_to_excel(data, file_path, title, include_chart)
                
            elif format_type == "Word":
                file_path = os.path.join(reports_dir, f"vacation_calendar_{timestamp}.docx")
                self._export_calendar_to_word(data, file_path, title, include_chart)
            
            logger.info(f"Календарный график отпусков успешно экспортирован в {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте календарного графика отпусков: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные: {str(e)}")
            return None

            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных о больничных: {str(e)}")
            messagebox.showerror("Ошибка", f"Не удалось экспортировать данные: {str(e)}")
            return None

    def _export_to_pdf(self, data, file_path, title, include_details=True, include_chart=True, is_summary=False):
        """
        Экспорт данных о зарплате в PDF формат
        
        :param data: Данные о зарплате
        :param file_path: Путь к создаваемому файлу
        :param title: Заголовок отчета
        :param include_details: Включать ли детализацию
        :param include_chart: Включать ли графики
        :param is_summary: Является ли отчет сводным
        """
        try:
            # Проверка наличия необходимых библиотек
            try:
                from reportlab.lib.pagesizes import A4
                from reportlab.lib import colors
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
                from reportlab.lib.styles import getSampleStyleSheet
                import matplotlib.pyplot as plt
                from io import BytesIO
            except ImportError:
                messagebox.showerror(
                    "Ошибка импорта", 
                    "Отсутствуют необходимые библиотеки для экспорта в PDF. Установите их командой:\n"
                    "pip install reportlab matplotlib"
                )
                return
            
            # Создание PDF документа
            doc = SimpleDocTemplate(file_path, pagesize=A4)
            styles = getSampleStyleSheet()
            elements = []
            
            # Заголовок отчета
            elements.append(Paragraph(title, styles['Title']))
            elements.append(Spacer(1, 12))
            
            # Дата создания отчета
            elements.append(Paragraph(f"Дата создания: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            if is_summary:
                # Сводная таблица для всех преподавателей
                table_data = [['Преподаватель', 'Валовая зарплата', 'Налоги', 'Чистая зарплата']]
                
                for teacher_data in data:
                    teacher_name = teacher_data.get('teacher_name', '')
                    gross_salary = f"{teacher_data.get('gross_salary', 0):.2f}"
                    tax_amount = f"{teacher_data.get('tax_amount', 0):.2f}"
                    net_salary = f"{teacher_data.get('net_salary', 0):.2f}"
                    
                    table_data.append([teacher_name, gross_salary, tax_amount, net_salary])
                
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(table)
                elements.append(Spacer(1, 12))
                
                # Добавление суммарных значений
                total_gross = sum(item.get('gross_salary', 0) for item in data)
                total_tax = sum(item.get('tax_amount', 0) for item in data)
                total_net = sum(item.get('net_salary', 0) for item in data)
                
                elements.append(Paragraph(f"Итого валовая зарплата: {total_gross:.2f} руб.", styles['Normal']))
                elements.append(Paragraph(f"Итого налоги: {total_tax:.2f} руб.", styles['Normal']))
                elements.append(Paragraph(f"Итого чистая зарплата: {total_net:.2f} руб.", styles['Normal']))
                
                # Если включены графики, добавляем круговую диаграмму распределения зарплат
                if include_chart:
                    plt.figure(figsize=(8, 6))
                    teacher_names = [item.get('teacher_name', '') for item in data]
                    net_salaries = [item.get('net_salary', 0) for item in data]
                    
                    plt.pie(net_salaries, labels=teacher_names, autopct='%1.1f%%', startangle=90)
                    plt.axis('equal')
                    plt.title('Распределение чистой зарплаты')
                    
                    buffer = BytesIO()
                    plt.savefig(buffer, format='png')
                    plt.close()
                    
                    # Добавляем диаграмму в отчет
                    buffer.seek(0)
                    img = Image(buffer, width=400, height=300)
                    elements.append(img)
            else:
                # Детальный отчет для одного преподавателя
                if not data:
                    elements.append(Paragraph("Нет данных для отображения", styles['Normal']))
                else:
                    # Основная информация о преподавателе
                    teacher_data = data[0] if isinstance(data, list) else data
                    
                    elements.append(Paragraph(f"Преподаватель: {teacher_data.get('teacher_name', '')}", styles['Heading2']))
                    elements.append(Paragraph(f"Должность: {teacher_data.get('position', '')}", styles['Normal']))
                    elements.append(Paragraph(f"Ставка: {teacher_data.get('hourly_rate', 0):.2f} руб./час", styles['Normal']))
                    elements.append(Spacer(1, 12))
                    
                    # Детали расчета зарплаты
                    if include_details:
                        elements.append(Paragraph("Детали расчета", styles['Heading3']))
                        
                        details_data = [
                            ['Параметр', 'Значение'],
                            ['Отработано часов', str(teacher_data.get('hours_worked', 0))],
                            ['Базовая зарплата', f"{teacher_data.get('base_salary', 0):.2f} руб."],
                            ['Надбавка за должность', f"{teacher_data.get('position_bonus', 0):.2f} руб."],
                            ['Надбавка за степень', f"{teacher_data.get('degree_bonus', 0):.2f} руб."],
                            ['Надбавка за стаж', f"{teacher_data.get('experience_bonus', 0):.2f} руб."],
                            ['Надбавка за категорию', f"{teacher_data.get('category_bonus', 0):.2f} руб."],
                            ['Надбавка молодому специалисту', f"{teacher_data.get('young_specialist_bonus', 0):.2f} руб."],
                            ['Оплата больничных', f"{teacher_data.get('sick_leave_pay', 0):.2f} руб."],
                            ['Бонус', f"{teacher_data.get('bonus', 0):.2f} руб."]
                        ]
                        
                        details_table = Table(details_data)
                        details_table.setStyle(TableStyle([
                            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                            ('GRID', (0, 0), (-1, -1), 1, colors.black)
                        ]))
                        
                        elements.append(details_table)
                        elements.append(Spacer(1, 12))
                    
                    # Итоговые суммы
                    elements.append(Paragraph("Итоговые суммы", styles['Heading3']))
                    
                    totals_data = [
                        ['Параметр', 'Значение'],
                        ['Валовая зарплата', f"{teacher_data.get('gross_salary', 0):.2f} руб."],
                        ['Ставка налога', f"{teacher_data.get('tax_rate', 0):.2f}%"],
                        ['Сумма налога', f"{teacher_data.get('tax_amount', 0):.2f} руб."],
                        ['Профсоюзные взносы', f"{teacher_data.get('union_contribution', 0):.2f} руб."],
                        ['ЧИСТАЯ ЗАРПЛАТА', f"{teacher_data.get('net_salary', 0):.2f} руб."]
                    ]
                    
                    totals_table = Table(totals_data)
                    totals_table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('BACKGROUND', (0, -1), (-1, -1), colors.lightblue),
                        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold')
                    ]))
                    
                    elements.append(totals_table)
                    
                    # Если включены графики, добавляем диаграмму структуры зарплаты
                    if include_chart:
                        plt.figure(figsize=(8, 6))
                        
                        # Компоненты зарплаты и их значения
                        components = [
                            'Базовая зарплата', 
                            'Надбавка за должность', 
                            'Надбавка за степень',
                            'Надбавка за стаж', 
                            'Надбавка за категорию', 
                            'Надбавка молодому специалисту',
                            'Оплата больничных', 
                            'Бонус'
                        ]
                        
                        values = [
                            teacher_data.get('base_salary', 0),
                            teacher_data.get('position_bonus', 0),
                            teacher_data.get('degree_bonus', 0),
                            teacher_data.get('experience_bonus', 0),
                            teacher_data.get('category_bonus', 0),
                            teacher_data.get('young_specialist_bonus', 0),
                            teacher_data.get('sick_leave_pay', 0),
                            teacher_data.get('bonus', 0)
                        ]
                        
                        # Фильтруем только ненулевые компоненты
                        non_zero_components = []
                        non_zero_values = []
                        
                        for comp, val in zip(components, values):
                            if val > 0:
                                non_zero_components.append(comp)
                                non_zero_values.append(val)
                        
                        if non_zero_values:
                            plt.pie(non_zero_values, labels=non_zero_components, autopct='%1.1f%%', startangle=90)
                            plt.axis('equal')
                            plt.title('Структура зарплаты')
                            
                            buffer = BytesIO()
                            plt.savefig(buffer, format='png')
                            plt.close()
                            
                            # Добавляем диаграмму в отчет
                            buffer.seek(0)
                            img = Image(buffer, width=400, height=300)
                            elements.append(img)
            
            # Создаем PDF документ
            doc.build(elements)
            
            logger.info(f"Данные о зарплате успешно экспортированы в PDF: {file_path}")
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных в PDF: {str(e)}")
            raise


    def _export_to_excel(self, data, file_path, title, include_details=True, include_chart=True, is_summary=False):
        """
        Экспорт данных о зарплате в Excel формат
        
        :param data: Данные о зарплате
        :param file_path: Путь к создаваемому файлу
        :param title: Заголовок отчета
        :param include_details: Включать ли детализацию
        :param include_chart: Включать ли графики
        :param is_summary: Является ли отчет сводным
        """
        try:
            # Проверка наличия необходимых библиотек
            try:
                import pandas as pd
                import matplotlib.pyplot as plt
                from io import BytesIO
                from openpyxl import Workbook
                from openpyxl.utils.dataframe import dataframe_to_rows
                from openpyxl.chart import PieChart, Reference
                from openpyxl.drawing.image import Image as XlsxImage
            except ImportError:
                messagebox.showerror(
                    "Ошибка импорта", 
                    "Отсутствуют необходимые библиотеки для экспорта в Excel. Установите их командой:\n"
                    "pip install pandas openpyxl matplotlib"
                )
                return
            
            # Создаем новую книгу Excel
            wb = Workbook()
            ws = wb.active
            ws.title = "Отчет о зарплате"
            
            # Добавляем заголовок
            ws['A1'] = title
            ws['A1'].font = ws['A1'].font.copy(size=14, bold=True)
            
            # Добавляем дату создания
            ws['A2'] = f"Дата создания: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}"
            
            current_row = 4  # Начальная строка для данных
            
            if is_summary:
                # Сводный отчет для всех преподавателей
                # Создаем DataFrame с данными
                df = pd.DataFrame(data)
                
                # Переименовываем столбцы для понятности
                columns_mapping = {
                    'teacher_name': 'Преподаватель',
                    'gross_salary': 'Валовая зарплата',
                    'tax_amount': 'Налоги',
                    'net_salary': 'Чистая зарплата'
                }
                
                df = df[list(columns_mapping.keys())].rename(columns=columns_mapping)
                
                # Добавляем данные в Excel
                for r_idx, row in enumerate(dataframe_to_rows(df, index=False, header=True), start=current_row):
                    for c_idx, value in enumerate(row, start=1):
                        cell = ws.cell(row=r_idx, column=c_idx, value=value)
                        # Форматируем числовые значения
                        if isinstance(value, (int, float)) and c_idx > 1:
                            cell.number_format = '#,##0.00'
                
                # Добавляем суммарную строку
                total_row = current_row + len(df) + 1
                ws.cell(row=total_row, column=1, value="ИТОГО")
                ws.cell(row=total_row, column=1).font = ws.cell(row=total_row, column=1).font.copy(bold=True)
                
                # Формулы суммирования для каждой числовой колонки
                for col in range(2, 5):
                    cell = ws.cell(row=total_row, column=col)
                    start_cell = ws.cell(row=current_row+1, column=col).coordinate
                    end_cell = ws.cell(row=total_row-1, column=col).coordinate
                    cell.value = f"=SUM({start_cell}:{end_cell})"
                    cell.number_format = '#,##0.00'
                    cell.font = cell.font.copy(bold=True)
                
                # Если нужны графики
                if include_chart and len(df) > 0:
                    # Добавляем лист для графиков
                    chart_sheet = wb.create_sheet(title="Графики")
                    
                    # Создаем круговую диаграмму распределения чистой зарплаты
                    pie = PieChart()
                    pie.title = "Распределение чистой зарплаты"
                    
                    # Копируем данные о преподавателях и зарплатах на лист с графиком
                    chart_sheet['A1'] = "Преподаватель"
                    chart_sheet['B1'] = "Чистая зарплата"
                    
                    for idx, row in enumerate(df.values, start=2):
                        chart_sheet.cell(row=idx, column=1, value=row[0])  # Имя преподавателя
                        chart_sheet.cell(row=idx, column=2, value=row[3])  # Чистая зарплата
                    
                    # Создаем ссылки на данные для графика
                    labels = Reference(chart_sheet, min_col=1, min_row=2, max_row=1+len(df))
                    data = Reference(chart_sheet, min_col=2, min_row=1, max_row=1+len(df))
                    
                    pie.add_data(data, titles_from_data=True)
                    pie.set_categories(labels)
                    
                    # Добавляем график на лист
                    chart_sheet.add_chart(pie, "D2")
            else:
                # Детальный отчет для одного преподавателя
                if not data:
                    ws['A4'] = "Нет данных для отображения"
                else:
                    teacher_data = data[0] if isinstance(data, list) else data
                    
                    # Основная информация о преподавателе
                    ws['A4'] = f"Преподаватель: {teacher_data.get('teacher_name', '')}"
                    ws['A5'] = f"Должность: {teacher_data.get('position', '')}"
                    ws['A6'] = f"Ставка: {teacher_data.get('hourly_rate', 0):.2f} руб./час"
                    
                    current_row = 8
                    
                    # Детали расчета зарплаты
                    if include_details:
                        ws.cell(row=current_row, column=1, value="Детали расчета")
                        ws.cell(row=current_row, column=1).font = ws.cell(row=current_row, column=1).font.copy(bold=True)
                        
                        current_row += 1
                        
                        # Заголовки таблицы деталей
                        ws.cell(row=current_row, column=1, value="Параметр")
                        ws.cell(row=current_row, column=2, value="Значение")
                        
                        # Форматирование заголовков
                        for col in range(1, 3):
                            cell = ws.cell(row=current_row, column=col)
                            cell.font = cell.font.copy(bold=True)
                        
                        # Добавляем детали расчета
                        details = [
                            ('Отработано часов', str(teacher_data.get('hours_worked', 0))),
                            ('Базовая зарплата', f"{teacher_data.get('base_salary', 0):.2f} руб."),
                            ('Надбавка за должность', f"{teacher_data.get('position_bonus', 0):.2f} руб."),
                            ('Надбавка за степень', f"{teacher_data.get('degree_bonus', 0):.2f} руб."),
                            ('Надбавка за стаж', f"{teacher_data.get('experience_bonus', 0):.2f} руб."),
                            ('Надбавка за категорию', f"{teacher_data.get('category_bonus', 0):.2f} руб."),
                            ('Надбавка молодому специалисту', f"{teacher_data.get('young_specialist_bonus', 0):.2f} руб."),
                            ('Оплата больничных', f"{teacher_data.get('sick_leave_pay', 0):.2f} руб."),
                            ('Бонус', f"{teacher_data.get('bonus', 0):.2f} руб.")
                        ]
                        
                        for detail in details:
                            current_row += 1
                            ws.cell(row=current_row, column=1, value=detail[0])
                            ws.cell(row=current_row, column=2, value=detail[1])
                        
                        current_row += 2
                    
                    # Итоговые суммы
                    ws.cell(row=current_row, column=1, value="Итоговые суммы")
                    ws.cell(row=current_row, column=1).font = ws.cell(row=current_row, column=1).font.copy(bold=True)
                    
                    current_row += 1
                    
                    # Заголовки таблицы итогов
                    ws.cell(row=current_row, column=1, value="Параметр")
                    ws.cell(row=current_row, column=2, value="Значение")
                    
                    # Форматирование заголовков
                    for col in range(1, 3):
                        cell = ws.cell(row=current_row, column=col)
                        cell.font = cell.font.copy(bold=True)
                    
                    # Добавляем итоговые значения
                    totals = [
                        ('Валовая зарплата', f"{teacher_data.get('gross_salary', 0):.2f} руб."),
                        ('Ставка налога', f"{teacher_data.get('tax_rate', 0):.2f}%"),
                        ('Сумма налога', f"{teacher_data.get('tax_amount', 0):.2f} руб."),
                        ('Профсоюзные взносы', f"{teacher_data.get('union_contribution', 0):.2f} руб."),
                        ('ЧИСТАЯ ЗАРПЛАТА', f"{teacher_data.get('net_salary', 0):.2f} руб.")
                    ]
                    
                    for i, total in enumerate(totals):
                        current_row += 1
                        ws.cell(row=current_row, column=1, value=total[0])
                        ws.cell(row=current_row, column=2, value=total[1])
                        
                        # Выделяем последнюю строку (чистая зарплата)
                        if i == len(totals) - 1:
                            for col in range(1, 3):
                                cell = ws.cell(row=current_row, column=col)
                                cell.font = cell.font.copy(bold=True)
                    
                    # Если нужны графики, добавляем диаграмму структуры зарплаты
                    if include_chart:
                        # Создаем лист для графиков
                        chart_sheet = wb.create_sheet(title="Структура зарплаты")
                        
                        # Компоненты зарплаты и их значения
                        components = [
                            'Базовая зарплата', 
                            'Надбавка за должность', 
                            'Надбавка за степень',
                            'Надбавка за стаж', 
                            'Надбавка за категорию', 
                            'Надбавка молодому специалисту',
                            'Оплата больничных', 
                            'Бонус'
                        ]
                        
                        values = [
                            teacher_data.get('base_salary', 0),
                            teacher_data.get('position_bonus', 0),
                            teacher_data.get('degree_bonus', 0),
                            teacher_data.get('experience_bonus', 0),
                            teacher_data.get('category_bonus', 0),
                            teacher_data.get('young_specialist_bonus', 0),
                            teacher_data.get('sick_leave_pay', 0),
                            teacher_data.get('bonus', 0)
                        ]
                        
                        # Фильтруем только ненулевые компоненты
                        non_zero_components = []
                        non_zero_values = []
                        
                        for comp, val in zip(components, values):
                            if val > 0:
                                non_zero_components.append(comp)
                                non_zero_values.append(val)
                        
                        if non_zero_values:
                            # Добавляем данные на лист
                            chart_sheet['A1'] = "Компонент зарплаты"
                            chart_sheet['B1'] = "Сумма (руб.)"
                            
                            for idx, (comp, val) in enumerate(zip(non_zero_components, non_zero_values), start=2):
                                chart_sheet.cell(row=idx, column=1, value=comp)
                                chart_sheet.cell(row=idx, column=2, value=val)
                            
                            # Создаем круговую диаграмму
                            pie = PieChart()
                            pie.title = "Структура зарплаты"
                            
                            # Создаем ссылки на данные для графика
                            labels = Reference(chart_sheet, min_col=1, min_row=2, max_row=1+len(non_zero_components))
                            data = Reference(chart_sheet, min_col=2, min_row=1, max_row=1+len(non_zero_components))
                            
                            pie.add_data(data, titles_from_data=True)
                            pie.set_categories(labels)
                            
                            # Добавляем график на лист
                            chart_sheet.add_chart(pie, "D2")
            
            # Сохраняем Excel-файл
            wb.save(file_path)
            
            logger.info(f"Данные о зарплате успешно экспортированы в Excel: {file_path}")
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных в Excel: {str(e)}")
            raise

    def _export_to_word(self, data, file_path, title, include_details=True, include_chart=True, is_summary=False):
        """
        Экспорт данных о зарплате в Word формат
        
        :param data: Данные о зарплате
        :param file_path: Путь к создаваемому файлу
        :param title: Заголовок отчета
        :param include_details: Включать ли детализацию
        :param include_chart: Включать ли графики
        :param is_summary: Является ли отчет сводным
        """
        try:
            # Проверка наличия необходимых библиотек
            try:
                from docx import Document
                from docx.shared import Pt, Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.enum.table import WD_TABLE_ALIGNMENT
                import matplotlib.pyplot as plt
                from io import BytesIO
            except ImportError:
                messagebox.showerror(
                    "Ошибка импорта", 
                    "Отсутствуют необходимые библиотеки для экспорта в Word. Установите их командой:\n"
                    "pip install python-docx matplotlib"
                )
                return
            
            # Создаем новый документ Word
            doc = Document()
            
            # Заголовок отчета
            doc.add_heading(title, level=1)
            
            # Дата создания отчета
            p = doc.add_paragraph(f"Дата создания: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}")
            
            if is_summary:
                # Сводный отчет для всех преподавателей
                doc.add_heading("Сводный отчет по зарплате", level=2)
                
                # Создаем таблицу
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                # Заголовки таблицы
                header_cells = table.rows[0].cells
                header_cells[0].text = "Преподаватель"
                header_cells[1].text = "Валовая зарплата"
                header_cells[2].text = "Налоги"
                header_cells[3].text = "Чистая зарплата"
                
                # Добавляем данные в таблицу
                for teacher_data in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = teacher_data.get('teacher_name', '')
                    row_cells[1].text = f"{teacher_data.get('gross_salary', 0):.2f}"
                    row_cells[2].text = f"{teacher_data.get('tax_amount', 0):.2f}"
                    row_cells[3].text = f"{teacher_data.get('net_salary', 0):.2f}"
                
                # Добавляем итоговые значения
                total_gross = sum(item.get('gross_salary', 0) for item in data)
                total_tax = sum(item.get('tax_amount', 0) for item in data)
                total_net = sum(item.get('net_salary', 0) for item in data)
                
                doc.add_paragraph(f"Итого валовая зарплата: {total_gross:.2f} руб.")
                doc.add_paragraph(f"Итого налоги: {total_tax:.2f} руб.")
                p = doc.add_paragraph(f"Итого чистая зарплата: {total_net:.2f} руб.")
                p.runs[0].bold = True
                
                # Если включены графики, добавляем круговую диаграмму распределения зарплат
                if include_chart and data:
                    doc.add_heading("Распределение чистой зарплаты", level=3)
                    
                    plt.figure(figsize=(8, 6))
                    teacher_names = [item.get('teacher_name', '') for item in data]
                    net_salaries = [item.get('net_salary', 0) for item in data]
                    
                    plt.pie(net_salaries, labels=teacher_names, autopct='%1.1f%%', startangle=90)
                    plt.axis('equal')
                    plt.title('Распределение чистой зарплаты')
                    
                    # Сохраняем график во временный файл
                    chart_path = os.path.join(os.path.dirname(file_path), "temp_chart.png")
                    plt.savefig(chart_path, format='png', dpi=300)
                    plt.close()
                    
                    # Добавляем изображение в документ
                    doc.add_picture(chart_path, width=Cm(15))
                    
                    # Удаляем временный файл
                    if os.path.exists(chart_path):
                        os.remove(chart_path)
            else:
                # Детальный отчет для одного преподавателя
                if not data:
                    doc.add_paragraph("Нет данных для отображения")
                else:
                    teacher_data = data[0] if isinstance(data, list) else data
                    
                    # Основная информация о преподавателе
                    doc.add_heading(f"Расчет зарплаты преподавателя: {teacher_data.get('teacher_name', '')}", level=2)
                    doc.add_paragraph(f"Должность: {teacher_data.get('position', '')}")
                    doc.add_paragraph(f"Ставка: {teacher_data.get('hourly_rate', 0):.2f} руб./час")
                    
                    # Детали расчета зарплаты
                    if include_details:
                        doc.add_heading("Детали расчета", level=3)
                        
                        # Создаем таблицу для деталей
                        details_table = doc.add_table(rows=1, cols=2)
                        details_table.style = 'Table Grid'
                        
                        # Заголовки таблицы
                        header_cells = details_table.rows[0].cells
                        header_cells[0].text = "Параметр"
                        header_cells[1].text = "Значение"
                        
                        # Добавляем данные в таблицу
                        details = [
                            ('Отработано часов', str(teacher_data.get('hours_worked', 0))),
                            ('Базовая зарплата', f"{teacher_data.get('base_salary', 0):.2f} руб."),
                            ('Надбавка за должность', f"{teacher_data.get('position_bonus', 0):.2f} руб."),
                            ('Надбавка за степень', f"{teacher_data.get('degree_bonus', 0):.2f} руб."),
                            ('Надбавка за стаж', f"{teacher_data.get('experience_bonus', 0):.2f} руб."),
                            ('Надбавка за категорию', f"{teacher_data.get('category_bonus', 0):.2f} руб."),
                            ('Надбавка молодому специалисту', f"{teacher_data.get('young_specialist_bonus', 0):.2f} руб."),
                            ('Оплата больничных', f"{teacher_data.get('sick_leave_pay', 0):.2f} руб."),
                            ('Бонус', f"{teacher_data.get('bonus', 0):.2f} руб.")
                        ]
                        
                        for param, value in details:
                            row_cells = details_table.add_row().cells
                            row_cells[0].text = param
                            row_cells[1].text = value
                    
                    # Итоговые суммы
                    doc.add_heading("Итоговые суммы", level=3)
                    
                    # Создаем таблицу для итогов
                    totals_table = doc.add_table(rows=1, cols=2)
                    totals_table.style = 'Table Grid'
                    
                    # Заголовки таблицы
                    header_cells = totals_table.rows[0].cells
                    header_cells[0].text = "Параметр"
                    header_cells[1].text = "Значение"
                    
                    # Добавляем данные в таблицу
                    totals = [
                        ('Валовая зарплата', f"{teacher_data.get('gross_salary', 0):.2f} руб."),
                        ('Ставка налога', f"{teacher_data.get('tax_rate', 0):.2f}%"),
                        ('Сумма налога', f"{teacher_data.get('tax_amount', 0):.2f} руб."),
                        ('Профсоюзные взносы', f"{teacher_data.get('union_contribution', 0):.2f} руб."),
                        ('ЧИСТАЯ ЗАРПЛАТА', f"{teacher_data.get('net_salary', 0):.2f} руб.")
                    ]
                    
                    for param, value in totals:
                        row_cells = totals_table.add_row().cells
                        row_cells[0].text = param
                        row_cells[1].text = value
                        
                        # Выделяем последнюю строку (чистая зарплата)
                        if param == 'ЧИСТАЯ ЗАРПЛАТА':
                            for cell in row_cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True
                    
                    # Если включены графики, добавляем диаграмму структуры зарплаты
                    if include_chart:
                        doc.add_heading("Структура зарплаты", level=3)
                        
                        # Компоненты зарплаты и их значения
                        components = [
                            'Базовая зарплата', 
                            'Надбавка за должность', 
                            'Надбавка за степень',
                            'Надбавка за стаж', 
                            'Надбавка за категорию', 
                            'Надбавка молодому специалисту',
                            'Оплата больничных', 
                            'Бонус'
                        ]
                        
                        values = [
                            teacher_data.get('base_salary', 0),
                            teacher_data.get('position_bonus', 0),
                            teacher_data.get('degree_bonus', 0),
                            teacher_data.get('experience_bonus', 0),
                            teacher_data.get('category_bonus', 0),
                            teacher_data.get('young_specialist_bonus', 0),
                            teacher_data.get('sick_leave_pay', 0),
                            teacher_data.get('bonus', 0)
                        ]
                        
                        # Фильтруем только ненулевые компоненты
                        non_zero_components = []
                        non_zero_values = []
                        
                        for comp, val in zip(components, values):
                            if val > 0:
                                non_zero_components.append(comp)
                                non_zero_values.append(val)
                        
                        if non_zero_values:
                            plt.figure(figsize=(8, 6))
                            plt.pie(non_zero_values, labels=non_zero_components, autopct='%1.1f%%', startangle=90)
                            plt.axis('equal')
                            plt.title('Структура зарплаты')
                            
                            # Сохраняем график во временный файл
                            chart_path = os.path.join(os.path.dirname(file_path), "temp_chart.png")
                            plt.savefig(chart_path, format='png', dpi=300)
                            plt.close()
                            
                            # Добавляем изображение в документ
                            doc.add_picture(chart_path, width=Cm(15))
                            
                            # Удаляем временный файл
                            if os.path.exists(chart_path):
                                os.remove(chart_path)
            
            # Сохраняем Word-файл
            doc.save(file_path)
            
            logger.info(f"Данные о зарплате успешно экспортированы в Word: {file_path}")
            
        except Exception as e:
            logger.error(f"Ошибка при экспорте данных в Word: {str(e)}")
            raise
    







    def _create_pdf_report(self, salary_data, teacher_info=None, period_start=None, period_end=None, output_file=None, include_chart=False):
        """
        Создает PDF-отчет с данными о зарплате
        
        Args:
            salary_data (list): Данные о зарплате
            teacher_info (dict, optional): Информация о преподавателе
            period_start (str, optional): Начало периода
            period_end (str, optional): Конец периода
            output_file (str, optional): Путь к выходному файлу
            include_chart (bool, optional): Включать ли график
        
        Returns:
            str: Путь к созданному файлу
        """
        try:
            # Проверяем наличие ReportLab
            try:
                from reportlab.pdfgen import canvas
                from reportlab.lib.pagesizes import A4, landscape
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
                from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
                from reportlab.lib import colors
                import reportlab.rl_config
                reportlab.rl_config.warnOnMissingFontGlyphs = 0  # Отключаем предупреждения о недостающих глифах
            except ImportError:
                logging.error("Не найден пакет reportlab для создания PDF")
                raise ValueError("Для создания PDF отчетов требуется установить пакет reportlab")
                
            # Запрашиваем файл для сохранения, если он не указан
            if not output_file:
                output_file = filedialog.asksaveasfilename(
                    title="Сохранить PDF-отчет о зарплате",
                    filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")],
                    defaultextension=".pdf"
                )
                
            if not output_file:
                return None
                
            # Регистрируем шрифт с поддержкой кириллицы
            try:
                # Путь к шрифту в зависимости от ОС
                if sys.platform.startswith('win'):
                    # Пути к стандартным шрифтам Windows с поддержкой кириллицы
                    fonts = [
                        ('Arial', 'C:/Windows/Fonts/arial.ttf'),
                        ('ArialBold', 'C:/Windows/Fonts/arialbd.ttf'),
                        ('Calibri', 'C:/Windows/Fonts/calibri.ttf'),
                        ('Verdana', 'C:/Windows/Fonts/verdana.ttf')
                    ]
                    
                    # Регистрируем первый доступный шрифт
                    font_registered = False
                    for font_name, font_path in fonts:
                        if os.path.exists(font_path):
                            pdfmetrics.registerFont(TTFont(font_name, font_path))
                            font_registered = True
                            self.pdf_font_name = font_name
                            logging.info(f"Зарегистрирован шрифт {font_name} из {font_path}")
                            break
                            
                    if not font_registered:
                        # Если ни один из стандартных шрифтов не найден, сообщаем об ошибке
                        raise FileNotFoundError("Не найдены шрифты с поддержкой кириллицы")
                else:
                    # Для Linux/Mac используем DejaVuSans или другие доступные шрифты
                    font_paths = [
                        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                        '/usr/share/fonts/TTF/DejaVuSans.ttf',
                        '/Library/Fonts/Arial Unicode.ttf'
                    ]
                    
                    font_registered = False
                    for font_path in font_paths:
                        if os.path.exists(font_path):
                            pdfmetrics.registerFont(TTFont('DejaVuSans', font_path))
                            font_registered = True
                            self.pdf_font_name = 'DejaVuSans'
                            logging.info(f"Зарегистрирован шрифт DejaVuSans из {font_path}")
                            break
                            
                    if not font_registered:
                        # Если шрифт не найден, сообщаем об ошибке
                        raise FileNotFoundError("Не найдены шрифты с поддержкой кириллицы")
                        
            except Exception as e:
                logging.error(f"Ошибка при регистрации шрифта: {str(e)}")
                # Пытаемся использовать встроенные шрифты ReportLab
                self.pdf_font_name = 'Helvetica'
                logging.info("Используем встроенный шрифт Helvetica (кириллица может не отображаться)")
                
            logging.info(f"Создание PDF-отчета о зарплате: {output_file}")
            
            # Создаем стили с новым шрифтом
            styles = getSampleStyleSheet()
            styles.add(ParagraphStyle(
                name='Normal_Cyrillic', 
                fontName=self.pdf_font_name,
                fontSize=10,
                leading=12
            ))
            styles.add(ParagraphStyle(
                name='Heading1_Cyrillic', 
                fontName=self.pdf_font_name,
                fontSize=16,
                leading=18,
                alignment=1  # По центру
            ))
            styles.add(ParagraphStyle(
                name='Heading2_Cyrillic', 
                fontName=self.pdf_font_name,
                fontSize=14,
                leading=16,
                alignment=1  # По центру
            ))
            styles.add(ParagraphStyle(
                name='Heading3_Cyrillic', 
                fontName=self.pdf_font_name,
                fontSize=12,
                leading=14,
                alignment=0  # По левому краю
            ))
            styles.add(ParagraphStyle(
                name='Footer_Cyrillic', 
                fontName=self.pdf_font_name,
                fontSize=8,
                leading=10,
                alignment=2  # По правому краю
            ))
            
            # Создаем PDF документ
            doc = SimpleDocTemplate(
                output_file, 
                pagesize=landscape(A4),
                title="Отчет по заработной плате",
                author="Система расчета зарплаты"
            )
            elements = []
            
            # Заголовок отчета
            elements.append(Paragraph("Отчет по заработной плате", styles['Heading1_Cyrillic']))
            elements.append(Spacer(1, 15))
            
            # Блок с информацией о преподавателе (если указан)
            if teacher_info:
                teacher_section = []
                teacher_section.append(Paragraph("Информация о преподавателе:", styles['Heading3_Cyrillic']))
                
                # Форматируем данные о преподавателе
                teacher_data = [
                    ["ФИО:", teacher_info.get('name', 'Не указано')],
                    ["Должность:", teacher_info.get('position', 'Не указано')],
                    ["Ученая степень:", teacher_info.get('academic_degree', 'Не указано')],
                    ["Квалификация:", teacher_info.get('qualification_category', 'Не указано')],
                    ["Стаж работы:", f"{teacher_info.get('experience_years', 'Не указано')} лет"],
                    ["Часовая ставка:", f"{teacher_info.get('hourly_rate', 'Не указано')} руб/ч"]
                ]
                
                # Отображаем дополнительные параметры, если они есть
                if 'hire_date' in teacher_info:
                    teacher_data.append(["Дата приема:", teacher_info.get('hire_date', 'Не указано')])
                
                if 'young_specialist' in teacher_info:
                    young_specialist = "Да" if teacher_info.get('young_specialist') else "Нет"
                    teacher_data.append(["Молодой специалист:", young_specialist])
                    
                if 'union_member' in teacher_info:
                    union_member = "Да" if teacher_info.get('union_member') else "Нет"
                    teacher_data.append(["Член профсоюза:", union_member])
                
                # Создаем таблицу с информацией о преподавателе
                teacher_table = Table(teacher_data, colWidths=[150, 300])
                teacher_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), self.pdf_font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                    ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.lightgrey),
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
                ]))
                
                elements.append(teacher_table)
                elements.append(Spacer(1, 15))
            
            # Информация о периоде отчета
            period_info = []
            if period_start or period_end:
                period_text = "Период отчета: "
                if period_start and period_end:
                    period_text += f"с {period_start} по {period_end}"
                elif period_start:
                    period_text += f"с {period_start}"
                elif period_end:
                    period_text += f"по {period_end}"
                    
                period_info.append(Paragraph(period_text, styles['Normal_Cyrillic']))
                period_info.append(Spacer(1, 10))
                
            elements.extend(period_info)
            
            # Проверяем, есть ли данные для отчета
            if not salary_data:
                elements.append(Paragraph("Нет данных для отображения в отчете", styles['Normal_Cyrillic']))
                elements.append(Spacer(1, 10))
            else:
                # Создаем заголовок таблицы данных
                elements.append(Paragraph("Данные о заработной плате:", styles['Heading3_Cyrillic']))
                elements.append(Spacer(1, 10))
                
                # Определяем ширину колонок для таблицы
                # Используем фиксированные значения вместо доступа к внутренним атрибутам
                col_widths = [100, 100, 100, 100, 100,]  # Ширины колонок
                
                # Заголовки таблицы
                data = [
                    ["Дата", "Отработано часов", 
                    "Надбавки", "Валовая зарплата", "Чистая зарплата"]
                ]
                
                # Функция для безопасного доступа к числовым полям
                def safe_get(row, field, default=0.0):
                    """Безопасно получает числовое значение из словаря, возвращая default в случае ошибки"""
                    try:
                        value = row.get(field)
                        if value is None:
                            return default
                        return float(value)
                    except (ValueError, TypeError):
                        return default
                
                # Данные о зарплате
                for row in salary_data:
                    # Обработка даты (может быть в разных форматах)
                    calc_date = row.get('calculation_date', '')
                    if hasattr(calc_date, 'strftime'):
                        formatted_date = calc_date.strftime('%d.%m.%Y')
                    else:
                        formatted_date = str(calc_date)
                    
                    # Форматирование числовых значений
                    def format_float(value, default="0.00"):
                        if value is None:
                            return default
                        try:
                            return f"{float(value):.2f}"
                        except (ValueError, TypeError):
                            return default
                    
                    
                    # Обрабатываем поле надбавок - если есть total_bonuses, используем его,
                    # иначе пытаемся сложить все отдельные бонусы
                    if 'total_bonuses' in row and row['total_bonuses'] is not None:
                        total_bonuses = safe_get(row, 'total_bonuses')
                    else:
                        total_bonuses = sum([
                            safe_get(row, 'position_bonus'),
                            safe_get(row, 'degree_bonus'),
                            safe_get(row, 'experience_bonus'),
                            safe_get(row, 'category_bonus'),
                            safe_get(row, 'young_specialist_bonus'),
                            safe_get(row, 'bonus')
                        ])
                    
                    gross_salary = safe_get(row, 'gross_salary')
                    tax_amount = safe_get(row, 'tax_amount')
                    union_contribution = safe_get(row, 'union_contribution')
                    net_salary = safe_get(row, 'net_salary')
                    
                    # Добавляем строку данных
                    data.append([
                        formatted_date,
                        str(row.get('hours_worked', '0')),
                        format_float(total_bonuses),
                        format_float(gross_salary),
                        format_float(net_salary)
                    ])
                
                # Создаем таблицу данных с указанными ширинами колонок
                table = Table(data, repeatRows=1, colWidths=col_widths)
                table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), self.pdf_font_name),
                    ('FONTSIZE', (0, 0), (-1, 0), 9),  # Размер шрифта для заголовков
                    ('FONTSIZE', (0, 1), (-1, -1), 8),  # Размер шрифта для данных
                    ('BACKGROUND', (0, 0), (-1, 0), colors.lightblue),  # Фон заголовков
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),  # Цвет текста заголовков
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),  # Выравнивание по центру
                    ('ALIGN', (1, 1), (1, -1), 'LEFT'),  # ФИО выравниваем по левому краю
                    ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),  # Вертикальное выравнивание
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),  # Сетка таблицы
                    ('BOX', (0, 0), (-1, -1), 1, colors.black),  # Рамка таблицы
                    ('LINEBELOW', (0, 0), (-1, 0), 2, colors.black),  # Линия под заголовками
                ]))
                
                elements.append(table)
                
                # Добавляем итоговую информацию
                elements.append(Spacer(1, 15))
                
                # Вычисляем итоги, используя безопасный доступ к полям
                
                total_hours = sum(safe_get(row, 'hours_worked') for row in salary_data)
                total_gross = sum(safe_get(row, 'gross_salary') for row in salary_data)
                total_net = sum(safe_get(row, 'net_salary') for row in salary_data)
                
                summary_data = [
                    ["Итого по отчету:", f"{total_hours:.2f}", f"{total_bonuses:.2f}", f"{total_gross:.2f}", f"{total_net:.2f}"]
                ]
                
                # Создаем таблицу итогов с теми же ширинами колонок
                summary_table = Table(summary_data, colWidths=col_widths)
                summary_table.setStyle(TableStyle([
                    ('FONTNAME', (0, 0), (-1, -1), self.pdf_font_name),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BACKGROUND', (0, 0), (-1, -1), colors.lightgrey),
                    ('ALIGN', (0, 0), (0, -1), 'RIGHT'),  # Выравнивание текста "Итого" по правому краю
                    ('ALIGN', (1, 0), (-1, -1), 'CENTER'),  # Выравнивание остальных ячеек по центру
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('SPAN', (0, 0), (2, 0)),  # Объединяем первые три ячейки для "Итого по отчету:"
                ]))
                
                elements.append(summary_table)
                
                # Если требуется график, добавляем его
                if include_chart:
                    try:
                        import matplotlib.pyplot as plt
                        import numpy as np
                        import io
                        
                        # Создаем временный график
                        plt.figure(figsize=(8, 4))
                        
                        # Данные для графика
                        dates = [row.get('calculation_date') for row in salary_data 
                                if hasattr(row.get('calculation_date', ''), 'strftime')]
                        dates = [d.strftime('%d.%m.%Y') for d in dates]
                        values = [safe_get(row, 'net_salary') for row in salary_data]
                        
                        if dates and values:
                            plt.bar(dates, values)
                            plt.title('Динамика заработной платы')
                            plt.xlabel('Дата')
                            plt.ylabel('Чистая зарплата (руб)')
                            plt.xticks(rotation=45)
                            plt.tight_layout()
                            
                            # Сохраняем график во временный буфер
                            buf = io.BytesIO()
                            plt.savefig(buf, format='png')
                            buf.seek(0)
                            
                            # Добавляем график в отчет
                            elements.append(Spacer(1, 20))
                            elements.append(Paragraph("Динамика заработной платы", styles['Heading3_Cyrillic']))
                            elements.append(Spacer(1, 10))
                            elements.append(Image(buf, width=500, height=250))
                            
                            plt.close()
                    except ImportError:
                        elements.append(Spacer(1, 10))
                        elements.append(Paragraph("Для отображения графиков требуется установить пакеты matplotlib и numpy", styles['Normal_Cyrillic']))
                    except Exception as chart_error:
                        logging.error(f"Ошибка при создании графика: {str(chart_error)}")
                        elements.append(Spacer(1, 10))
                        elements.append(Paragraph(f"Не удалось создать график: {str(chart_error)}", styles['Normal_Cyrillic']))
            
            # Добавляем дату создания отчета и подпись
            elements.append(Spacer(1, 20))
            current_date = datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")
            elements.append(Paragraph(f"Отчет создан: {current_date}", styles['Footer_Cyrillic']))
            elements.append(Paragraph("Система расчета зарплаты для лицея ПолессГУ", styles['Footer_Cyrillic']))
            
            # Строим документ
            doc.build(elements)
            
            logging.info(f"PDF-отчет успешно создан: {output_file}")
            return output_file
            
        except Exception as e:
            logging.error(f"Ошибка при создании PDF-отчета: {str(e)}")
            raise ValueError(f"Не удалось создать PDF-отчет: {str(e)}")


